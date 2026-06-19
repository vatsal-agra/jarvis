#!/usr/bin/env python
"""
Jarvis — AI Voice Assistant
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Brain      : Google Gemini 2.5 (free tier, multi-key + multi-model rotation)
             fully cloud — nothing runs on the local GPU
Vision     : Gemini multimodal (sees screen + browser)
Web        : Playwright   persistent Chrome (DOM-based, 99% reliable)
Search     : DuckDuckGo   (free, no key)
Native apps: pywinauto + pyautogui (Windows accessibility + desktop control)
TTS        : edge-tts     Microsoft Neural voices
STT        : SpeechRecognition + Google API

Install:
    pip install SpeechRecognition pyaudio edge-tts pygame wikipedia
              deep-translator pyshorteners requests ddgs playwright pywinauto
              keyboard psutil pyperclip pypdf python-docx Pillow pyautogui
    python -m playwright install chromium
Keys: put JARVIS_GEMINI_KEYS=key1,key2,... in a .env file (or jarvis_keys.txt).
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""

import asyncio
import base64
import datetime
import io
import json
import math
import os
import random as _random
import re
import subprocess
import tempfile
import threading
import time
import urllib.parse
import webbrowser
from datetime import date

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  INTERRUPT / BARGE-IN
#  Press ESC at any time to abort the task Jarvis is currently doing (mid tool
#  loop and mid-speech) and return immediately to listening. This is what lets
#  you stop a wrong task the moment you notice it misheard you.
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
_ABORT = threading.Event()
_SPEAK_LOCK = threading.Lock()   # serialize TTS so voice/reminders/presence never overlap


def _install_abort_hotkey() -> str:
    """Make ESC abort the current task. We arm TWO independent listeners so at
    least one works in any environment:

      1. A global low-level hotkey via the `keyboard` lib. This fires even when
         Jarvis isn't the focused window — essential, because during browser
         automation Chrome is in front, not the terminal. (May need admin on
         some systems / not fire over RDP, hence the second listener.)
      2. A console watcher: on Windows, msvcrt reads single keystrokes, so a
         bare ESC press in the terminal works with no admin and no Enter. On
         other OSes (or if msvcrt is missing) it degrades to reading a line,
         so pressing Enter aborts.

    Both just call _ABORT.set(), which is idempotent — double-firing is fine."""
    arms = []

    # 1) global hotkey (cross-window)
    try:
        import keyboard
        keyboard.add_hotkey("esc", _ABORT.set)
        arms.append("global")
    except Exception:
        pass

    # 2) console watcher (no admin needed)
    try:
        import msvcrt

        def _console_watch():
            while True:
                try:
                    ch = msvcrt.getwch()          # blocks for one keystroke
                except Exception:
                    break
                if ch == "\x1b":                  # ESC
                    _ABORT.set()
        threading.Thread(target=_console_watch, daemon=True).start()
        arms.append("console-esc")
    except Exception:
        def _stdin_watch():
            while True:
                try:
                    input()
                except (EOFError, OSError):
                    break
                _ABORT.set()
        threading.Thread(target=_stdin_watch, daemon=True).start()
        arms.append("stdin-enter")

    if "stdin-enter" in arms and "global" not in arms:
        return "Press Enter in this window any time to stop the current task."
    return "Press ESC any time to stop the current task."

try:
    import ollama          # optional — only used if you re-enable a local fallback
except Exception:
    ollama = None
import requests
import speech_recognition as sr


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  HUD  ─  live web control center
#  A tiny stdlib HTTP server streams the assistant's state (listening,
#  thinking, which brain/key, every tool call, the reply, Gemini quota) to a
#  browser dashboard (jarvis_hud.html) over Server-Sent Events. Purely
#  observational — if anything here fails, the assistant runs exactly as before.
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
import http.server
import queue as _queue
import socketserver

_HUD_PORT          = int(os.environ.get("JARVIS_HUD_PORT", "8765"))
_HUD_DIR           = os.path.dirname(os.path.abspath(__file__))
_HUD_CLIENTS       = []                 # one queue.Queue per connected browser
_HUD_LOCK          = threading.Lock()
_HUD_LAST          = {}                 # last value of replayable events (for late joiners)
_GEMINI_USAGE      = {}                 # (key_index, model) -> requests made this session
_TEXT_CMDS         = _queue.Queue()     # commands typed into the HUD → main loop


def _hud_emit(kind: str, **data) -> None:
    """Broadcast one event to every connected HUD browser. Never raises."""
    try:
        data["kind"] = kind
        if kind in ("state", "brain", "quota", "online"):
            _HUD_LAST[kind] = data          # remember so a fresh page isn't blank
        msg = json.dumps(data)
        with _HUD_LOCK:
            for q in list(_HUD_CLIENTS):
                try:
                    q.put_nowait(msg)
                except Exception:
                    pass
    except Exception:
        pass


def _hud_quota() -> None:
    """Emit the per-key/per-model request meter (429 marks a key full)."""
    rows = []
    for (idx, model), used in sorted(_GEMINI_USAGE.items()):
        rows.append({"key": idx + 1, "model": model, "used": min(used, 20), "limit": 20})
    if rows:
        _hud_emit("quota", usage=rows)


# ── Quota persistence ─────────────────────────────────────────────────────────
# Google resets the free daily quota at midnight US-Pacific, so usage is keyed to
# the Pacific date and saved to disk — the meter survives restarts and only zeroes
# out when a genuinely new quota day begins.
_USAGE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "jarvis_usage.json")
_USAGE_DATE = None
_USAGE_LOCK = threading.Lock()   # guards the quota counter + jarvis_usage.json writes
                                 # (proactive vision loops hit Gemini outside _BRAIN_LOCK)


def _pacific_date() -> str:
    try:
        from zoneinfo import ZoneInfo
        return datetime.datetime.now(ZoneInfo("America/Los_Angeles")).strftime("%Y-%m-%d")
    except Exception:
        return (datetime.datetime.utcnow() - datetime.timedelta(hours=8)).strftime("%Y-%m-%d")


def _usage_save() -> None:
    try:
        with open(_USAGE_PATH, "w", encoding="utf-8") as f:
            json.dump({"date": _USAGE_DATE or _pacific_date(),
                       "usage": {f"{i}|{mdl}": v for (i, mdl), v in _GEMINI_USAGE.items()}}, f)
    except Exception:
        pass


def _usage_load() -> None:
    global _GEMINI_USAGE, _USAGE_DATE
    _USAGE_DATE = _pacific_date()
    try:
        with open(_USAGE_PATH, encoding="utf-8") as f:
            d = json.load(f)
        if d.get("date") == _USAGE_DATE:          # same quota day → restore counts
            _GEMINI_USAGE = {(int(k.split("|", 1)[0]), k.split("|", 1)[1]): v
                             for k, v in d.get("usage", {}).items()}
    except Exception:
        pass


def _usage_rollover() -> None:
    """If the Pacific day has changed, the quota reset — clear the meter."""
    global _USAGE_DATE
    today = _pacific_date()
    if _USAGE_DATE != today:
        with _USAGE_LOCK:
            _USAGE_DATE = today
            _GEMINI_USAGE.clear()
            _usage_save()


def _tool_summary(name: str, args: dict) -> str:
    """One-line human summary of a tool call for the action feed."""
    a = args or {}
    for k in ("query", "url", "text", "name", "command", "element", "content", "key"):
        if a.get(k):
            v = str(a[k]).replace("\n", " ")
            if k in ("element", "content"):
                v = f"[{a.get('element','')}] {str(a.get('content',''))}".strip()
            return v[:60]
    return ", ".join(f"{k}={v}" for k, v in list(a.items())[:2])[:60]


def _tool_outcome(result) -> tuple:
    """Map a tool result string to (status, short-detail) for the HUD."""
    s = str(result)
    low = s.lower()
    bad = any(t in low for t in ("error", "could not", "no element", "couldn't", "unavailable", "failed"))
    first = s.strip().splitlines()[0] if s.strip() else ""
    return ("err" if bad else "ok"), first[:54]


class _HudHandler(http.server.BaseHTTPRequestHandler):
    def log_message(self, *a):
        pass

    def do_GET(self):
        if self.path in ("/", "/index.html"):
            try:
                with open(os.path.join(_HUD_DIR, "jarvis_hud.html"), "rb") as f:
                    body = f.read()
            except Exception:
                body = b"<h1 style='color:#fff;background:#000'>jarvis_hud.html not found</h1>"
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            self.wfile.write(body)
            return

        if self.path == "/events":
            self.send_response(200)
            self.send_header("Content-Type", "text/event-stream")
            self.send_header("Cache-Control", "no-cache")
            self.send_header("Connection", "keep-alive")
            self.send_header("Access-Control-Allow-Origin", "*")
            self.end_headers()
            q = _queue.Queue(maxsize=256)
            with _HUD_LOCK:
                _HUD_CLIENTS.append(q)
            try:
                self.wfile.write(b": connected\n\n")
                for k in ("online", "brain", "quota", "state"):   # replay current state
                    if k in _HUD_LAST:
                        self.wfile.write(f"data: {json.dumps(_HUD_LAST[k])}\n\n".encode())
                self.wfile.flush()
                while True:
                    try:
                        msg = q.get(timeout=15)
                    except _queue.Empty:
                        self.wfile.write(b": ping\n\n")    # keep-alive heartbeat
                        self.wfile.flush()
                        continue
                    self.wfile.write(f"data: {msg}\n\n".encode())
                    self.wfile.flush()
            except Exception:
                pass
            finally:
                with _HUD_LOCK:
                    try:
                        _HUD_CLIENTS.remove(q)
                    except ValueError:
                        pass
            return

        self.send_response(404)
        self.end_headers()

    def do_POST(self):
        if self.path == "/command":
            try:
                n = int(self.headers.get("Content-Length", 0))
                body = json.loads(self.rfile.read(n) or b"{}")
                text = (body.get("text") or "").strip()
                if text:
                    _TEXT_CMDS.put(text)
            except Exception:
                pass
            self.send_response(204)
            self.send_header("Access-Control-Allow-Origin", "*")
            self.end_headers()
            return
        self.send_response(404)
        self.end_headers()


def _start_hud() -> str:
    """Launch the HUD server in a background thread. Returns its URL, or '' on failure."""
    try:
        socketserver.ThreadingTCPServer.allow_reuse_address = True
        srv = socketserver.ThreadingTCPServer(("127.0.0.1", _HUD_PORT), _HudHandler)
        srv.daemon_threads = True
        threading.Thread(target=srv.serve_forever, daemon=True).start()
        return f"http://127.0.0.1:{_HUD_PORT}"
    except Exception:
        return ""


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  TTS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def _clean_for_speech(text: str) -> str:
    text = re.sub(r"\*+", "", text)
    text = re.sub(r"#+\s*", "", text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    text = re.sub(r"`+", "", text)
    text = re.sub(r"-{2,}", "—", text)
    return text.strip()


def _speak_offline(text: str) -> bool:
    """OFFLINE fallback voice — Windows' built-in SAPI via PowerShell. No network,
    no extra deps, and it runs in its own process so it can't clash with pygame's
    COM apartment. Used when edge-tts can't reach Microsoft's servers."""
    try:
        safe = text.replace("'", "''")
        subprocess.run(
            ["powershell", "-NoProfile", "-Command",
             "Add-Type -AssemblyName System.Speech; "
             "$s = New-Object System.Speech.Synthesis.SpeechSynthesizer; "
             f"$s.Rate = 1; $s.Speak('{safe}')"],
            timeout=60, creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
        )
        return True
    except Exception as e:
        print(f"[Offline TTS error] {e}")
        return False


_edge_tts_broken = False   # once edge-tts fails (no DNS), skip it and go straight offline

try:
    import edge_tts
    import pygame

    VOICE = "en-US-GuyNeural"
    pygame.mixer.pre_init(44100, -16, 2, 512)
    pygame.mixer.init()

    async def _tts_async(text: str) -> str:
        communicate = edge_tts.Communicate(text, voice=VOICE)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
        tmp.close()
        await communicate.save(tmp.name)
        return tmp.name

    def speak(text: str) -> None:
        global _edge_tts_broken
        text = _clean_for_speech(str(text))
        print(f"\n🤖 Jarvis: {text}\n")
        _hud_emit("jarvis", text=text)
        _hud_emit("state", state="speaking", sub="responding")
        with _SPEAK_LOCK:                      # never let two voices overlap
            spoke = False
            if not _edge_tts_broken:
                try:
                    import concurrent.futures
                    # Fresh thread so asyncio.run() never conflicts with Playwright's loop.
                    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
                        fname = pool.submit(lambda: asyncio.run(_tts_async(text))).result(timeout=30)
                    pygame.mixer.music.load(fname)
                    pygame.mixer.music.play()
                    _spk = 0
                    while pygame.mixer.music.get_busy():
                        if _ABORT.is_set():           # ESC pressed → cut speech short
                            pygame.mixer.music.stop()
                            break
                        _spk += 1
                        env = 0.45 + 0.35 * abs(math.sin(_spk * 0.6)) + 0.15 * abs(math.sin(_spk * 1.9))
                        _hud_emit("level", v=min(1.0, env))
                        pygame.time.wait(60)
                    pygame.mixer.music.unload()
                    os.unlink(fname)
                    spoke = True
                except Exception as e:
                    print(f"[edge-tts unavailable: {e}]\n   → switching to offline Windows voice for the rest of this session.")
                    _edge_tts_broken = True
            if not spoke:
                _speak_offline(text)
        _hud_emit("level", v=0.0)
        _hud_emit("state", state="standby", sub="ready")

except ImportError:
    import pyttsx3

    _engine = pyttsx3.init("sapi5")
    _voices = _engine.getProperty("voices")
    _engine.setProperty("voice", _voices[0].id)
    _engine.setProperty("rate", 170)

    def speak(text: str) -> None:
        text = _clean_for_speech(str(text))
        print(f"\n🤖 Jarvis: {text}\n")
        _hud_emit("jarvis", text=text)
        _hud_emit("state", state="speaking", sub="responding")
        with _SPEAK_LOCK:
            _engine.say(text)
            _engine.runAndWait()
        _hud_emit("state", state="standby", sub="ready")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  STT  —  Silero VAD + Google Speech Recognition
#
#  How it works:
#   • Records in 32 ms chunks, feeding each to Silero VAD (neural model)
#   • Starts accumulating only when VAD confirms speech has begun
#   • Ends recording only after 1.5 s of *consistently* low speech probability
#   • Mid-sentence pauses (thinking, breathing) reset the silence counter
#     the moment any speech resumes — so it never cuts you off
#   • Finished audio is sent to Google STT
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
import numpy as np
# NOTE: torch is the slow import (~9s) and is used ONLY for the Silero voice
# detector, so we DON'T import it at module load. It's imported lazily by
# _get_vad() and warmed in a background thread at startup — Jarvis boots in ~2s
# and the model is ready by the time you first speak. No quality change.

_recognizer = sr.Recognizer()   # kept for Google STT only
_vad_model   = None             # loaded once at first use
_torch       = None             # the torch module, imported lazily with the VAD


def _get_vad():
    global _vad_model, _torch
    if _vad_model is None:
        import torch as _t
        _torch = _t
        from silero_vad import load_silero_vad
        _vad_model = load_silero_vad()
    return _vad_model


_mic_listed = False


def _list_input_devices(pa) -> None:
    """Print available microphones once, so a wrong default device is obvious."""
    global _mic_listed
    if _mic_listed:
        return
    _mic_listed = True
    try:
        default = pa.get_default_input_device_info()
        print(f"🎤 Mic: using '{default['name']}' (index {default['index']}). "
              f"Override with JARVIS_MIC_INDEX if this is wrong. Inputs:")
        for i in range(pa.get_device_count()):
            d = pa.get_device_info_by_index(i)
            if d.get("maxInputChannels", 0) > 0:
                print(f"     [{i}] {d['name']}")
    except Exception as e:
        print(f"🎤 Mic: no default input device found ({e}) — check your microphone.")


def takeCommand() -> str:
    import pyaudio as _pa

    RATE             = 16000   # Hz  — Silero requires 16 kHz
    CHUNK            = 512     # samples per chunk = 32 ms
    START_THRESH     = 0.4     # VAD prob to consider speech started
    END_THRESH       = 0.3     # VAD prob below which silence counter ticks
    END_CHUNKS_NEED  = 47      # 47 × 32 ms ≈ 1.5 s of silence → done
    MIN_SPEECH_START = 2       # consecutive speech chunks needed to begin recording
    PRE_ROLL         = 16      # chunks of audio kept before speech starts (~0.5 s)
    MAX_CHUNKS       = 625     # hard cap: 20 s
    WAIT_TIMEOUT     = 312     # give up waiting for speech after ~10 s

    model = _get_vad()
    model.reset_states()       # clear RNN state from last call

    pa = _pa.PyAudio()
    _list_input_devices(pa)
    open_kw = {"rate": RATE, "channels": 1, "format": _pa.paInt16,
               "input": True, "frames_per_buffer": CHUNK}
    # Prefer selecting the mic by NAME (stable) — device INDEXES shift whenever
    # you plug/unplug audio gear (e.g. a webcam adds mic entries and renumbers
    # everything), which silently breaks a pinned JARVIS_MIC_INDEX.
    mic_name = os.environ.get("JARVIS_MIC_NAME", "").strip().lower()
    mic_index = os.environ.get("JARVIS_MIC_INDEX", "").strip()
    chosen = None
    if mic_name:
        try:
            for i in range(pa.get_device_count()):
                d = pa.get_device_info_by_index(i)
                if d.get("maxInputChannels", 0) > 0 and mic_name in d.get("name", "").lower():
                    chosen = i
                    break
        except Exception:
            pass
        if chosen is None:
            print(f"🎤 No input device matching name '{mic_name}' — using the system default.")
    if chosen is None and mic_index.isdigit():
        chosen = int(mic_index)
    if chosen is not None:
        open_kw["input_device_index"] = chosen
        try:
            print(f"🎤 Listening on device [{chosen}] {pa.get_device_info_by_index(chosen)['name']}")
        except Exception:
            pass
    try:
        stream = pa.open(**open_kw)
    except Exception as e:
        print(f"🎤 Could not open the microphone: {e}")
        pa.terminate()
        time.sleep(1)
        return "none"

    print("🎙  Listening...")
    _hud_emit("state", state="listening", sub="listening for “Jarvis…”")

    pre_roll       = []        # ring buffer before speech starts
    recording      = []        # accumulated speech frames
    speech_started = False
    consec_speech  = 0         # consecutive chunks above START_THRESH
    silence_count  = 0         # consecutive chunks below END_THRESH
    total_waited   = 0
    lvl_tick       = 0         # throttles HUD audio-level emits
    peak_prob      = 0.0       # diagnostics: loudest speech-probability this cycle
    peak_rms       = 0.0       # diagnostics: loudest raw level this cycle

    try:
        while True:
            raw     = stream.read(CHUNK, exception_on_overflow=False)
            samples = np.frombuffer(raw, dtype=np.int16).astype(np.float32) / 32768.0
            prob    = model(_torch.from_numpy(samples), RATE).item()

            rms = float(np.sqrt(np.mean(samples * samples)))
            peak_prob = max(peak_prob, prob)
            peak_rms = max(peak_rms, rms)
            # Stream the live mic amplitude to the HUD reactor (every ~64 ms).
            lvl_tick += 1
            if lvl_tick % 2 == 0:
                _hud_emit("level", v=min(1.0, rms * 7.0))

            if not speech_started:
                # Keep a short pre-roll buffer so we don't clip the first syllable
                pre_roll.append(raw)
                if len(pre_roll) > PRE_ROLL:
                    pre_roll.pop(0)

                if prob >= START_THRESH:
                    consec_speech += 1
                else:
                    consec_speech = 0

                if consec_speech >= MIN_SPEECH_START:
                    speech_started = True
                    recording      = list(pre_roll)   # include pre-roll
                    silence_count  = 0
                    print("   Recording...")
                    _hud_emit("state", state="listening", sub="hearing you…")

                total_waited += 1
                if total_waited >= WAIT_TIMEOUT:
                    break   # nothing heard for 10 s

            else:
                recording.append(raw)

                if prob >= END_THRESH:
                    silence_count = 0   # speech resumed — reset counter
                else:
                    silence_count += 1

                if silence_count >= END_CHUNKS_NEED:
                    print("   End of speech detected.")
                    break

                if len(recording) >= MAX_CHUNKS:
                    break

    finally:
        stream.stop_stream()
        stream.close()
        pa.terminate()

    if not recording or not speech_started:
        # Diagnostics: did the mic capture ANY sound this cycle?
        if peak_rms < 0.005:
            print(f"   (heard silence — mic level ~0. Peak level {peak_rms:.4f}. "
                  f"Check the mic isn't muted / is the right device — see JARVIS_MIC_INDEX.)")
        else:
            print(f"   (no speech detected — peak level {peak_rms:.3f}, "
                  f"peak speech-prob {peak_prob:.2f}. Try speaking a bit louder/closer.)")
        return "none"

    # Hand raw PCM bytes to Google STT via SpeechRecognition
    audio_obj = sr.AudioData(b"".join(recording), RATE, 2)
    try:
        print("🔍 Recognizing...")
        query = _recognizer.recognize_google(audio_obj, language="en-in")
        print(f"👤 You: {query}")
        return query
    except sr.UnknownValueError:
        return "none"
    except sr.RequestError:
        speak("Speech service unavailable. Check your internet.")
        return "none"


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  BROWSER STATE
#  Priority:
#  Strategy:
#    1. If Jarvis Chrome is already running (port 9223) → reconnect via CDP
#    2. Otherwise → launch Chrome normally via subprocess (zero automation flags)
#       then connect to it via CDP — gives a fully normal, interactable window
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
_pw_instance  = None
_pw_browser   = None   # CDP-connected browser object
_pw_page      = None
_pw_known_pages = 0    # how many tabs we last saw — to detect a NEW popup tab
_jarvis_proc  = None   # subprocess handle for the Jarvis Chrome process

JARVIS_CHROME_PROFILE = os.path.join(os.path.expanduser("~"), ".jarvis_chrome_profile")
JARVIS_CDP_PORT       = 9223   # separate port so we never clash with user's Chrome

_CHROME_PATHS = [
    r"C:\Program Files\Google\Chrome\Application\chrome.exe",
    r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
    os.path.join(os.path.expanduser("~"),
                 r"AppData\Local\Google\Chrome\Application\chrome.exe"),
]


def _find_chrome() -> str | None:
    for p in _CHROME_PATHS:
        if os.path.exists(p):
            return p
    return None


def _ensure_playwright():
    global _pw_instance
    if _pw_instance is None:
        from playwright.sync_api import sync_playwright
        _pw_instance = sync_playwright().start()


def _cdp_page(port: int):
    """Connect to Chrome at *port* via CDP and return its first page."""
    try:
        requests.get(f"http://localhost:{port}/json/version", timeout=1)
        browser = _pw_instance.chromium.connect_over_cdp(f"http://localhost:{port}")
        ctx = browser.contexts[0] if browser.contexts else None
        if ctx:
            return browser, (ctx.pages[0] if ctx.pages else ctx.new_page())
    except Exception:
        pass
    return None, None


def _launch_jarvis_chrome():
    """
    Two-phase Chrome launch:

    Phase 1 (first run only) — login setup
      Launch Chrome with NO remote-debugging-port and NO CDP connection at all.
      The window is 100% normal: Google OAuth, LinkedIn login, everything works.
      User logs in, then closes Chrome.  We wait for it to exit.

    Phase 2 — automation
      Relaunch Chrome with --remote-debugging-port=JARVIS_CDP_PORT.
      Playwright connects via CDP.  User is already logged in so no OAuth flows
      happen and the debugger-pause trap is never triggered.
    """
    global _pw_browser, _jarvis_proc

    os.makedirs(JARVIS_CHROME_PROFILE, exist_ok=True)
    first_launch = not os.path.exists(
        os.path.join(JARVIS_CHROME_PROFILE, "Default", "Preferences")
    )

    chrome_exe = _find_chrome()
    if not chrome_exe:
        print("[Browser] Chrome not found — falling back to Playwright Chromium.")
        _pw_browser = _pw_instance.chromium.launch_persistent_context(
            JARVIS_CHROME_PROFILE, headless=False, slow_mo=80,
        )
        page = _pw_browser.pages[0] if _pw_browser.pages else _pw_browser.new_page()
        return page

    # ── Phase 1: login setup (first run only) ─────────────────────────────────
    if first_launch:
        print("\n" + "="*60)
        print("  FIRST-TIME BROWSER SETUP")
        print("  A normal Chrome window will open.")
        print("  Log into Google, LinkedIn, YouTube — any site Jarvis will use.")
        print("  When you are done, CLOSE that Chrome window.")
        print("  Jarvis will continue automatically.")
        print("="*60 + "\n")
        speak(
            "First time setup. A Chrome window is opening with no automation. "
            "Please log into your accounts, then close Chrome when you are done. "
            "Jarvis will continue automatically."
        )
        # Launch completely plain Chrome — zero debug port, zero CDP
        setup_proc = subprocess.Popen([
            chrome_exe,
            f"--user-data-dir={JARVIS_CHROME_PROFILE}",
            "--no-first-run",
            "--no-default-browser-check",
            "--start-maximized",
        ])
        print("  ⏳ Waiting for you to finish logging in and close Chrome...")
        setup_proc.wait()   # blocks until the user closes Chrome
        print("  ✅ Login saved. Launching automation browser...\n")
        speak("Logins saved. Starting up now.")

    # ── Phase 2: automation launch ─────────────────────────────────────────────
    _jarvis_proc = subprocess.Popen([
        chrome_exe,
        f"--user-data-dir={JARVIS_CHROME_PROFILE}",
        f"--remote-debugging-port={JARVIS_CDP_PORT}",
        "--no-first-run",
        "--no-default-browser-check",
        "--start-maximized",
    ])
    # Wait up to 10 s for the CDP endpoint to become available
    for _ in range(20):
        time.sleep(0.5)
        browser, page = _cdp_page(JARVIS_CDP_PORT)
        if page:
            _pw_browser = browser
            return page

    print("[Browser] Timed out waiting for Chrome CDP — is Chrome installed?")
    return None


def _get_page():
    """Return active Playwright page, launching Jarvis Chrome if needed."""
    global _pw_page, _pw_browser, _pw_known_pages
    _ensure_playwright()

    # Check if existing page is still alive
    if _pw_page is not None:
        try:
            closed = _pw_page.is_closed()
        except Exception:
            closed = True   # CDP connection broken = Chrome was closed
        if closed:
            _pw_page = None  # discard stale reference so relaunch triggers below
        else:
            # Follow a popup ONLY when a NEW tab actually appeared (count grew).
            # Don't blindly jump to the last tab every call — that used to silently
            # undo an explicit browser_tab switch.
            if _pw_browser:
                try:
                    live = [p for ctx in _pw_browser.contexts for p in ctx.pages
                            if not p.is_closed()]
                    if live and len(live) > _pw_known_pages:
                        _pw_page = live[-1]      # a popup/new tab opened → follow it
                    _pw_known_pages = len(live)
                except Exception:
                    pass
            return _pw_page

    # Try reconnecting to Jarvis Chrome if it's still running
    browser, page = _cdp_page(JARVIS_CDP_PORT)
    if page:
        _pw_browser = browser
        _pw_page = page
        print("[Browser] Reconnected to Jarvis Chrome via CDP.")
        return _pw_page

    # Chrome not running — launch it fresh
    _pw_page = _launch_jarvis_chrome()
    return _pw_page


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  INDEXED-ELEMENT SNAPSHOT
#
#  The reliable way to drive an LLM browser agent: instead of asking the model
#  to fuzzy-match element text (which breaks constantly), we tag EVERY visible
#  interactive element with a number (data-jarvis-id) and hand the model a clean
#  numbered list. The model just says "click 5" or "type into 3" — no guessing.
#  This is site-agnostic: it works on any page because it reads the live DOM.
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
_INDEX_JS = r"""
() => {
  const txt = (el) => (el && el.innerText || '').replace(/\s+/g, ' ').trim();

  // Resolve the best human-readable name for an element. Crucially this follows
  // <label for=id>, wrapping <label>, and aria-labelledby — so standard form
  // fields (GitHub's "Repository name", etc.) get a real name instead of blank.
  const nameOf = (el, isInput) => {
    let t = (el.getAttribute('aria-label') || '').trim();
    if (t) return t;
    const lb = el.getAttribute('aria-labelledby');
    if (lb) {
      const s = lb.split(/\s+/).map(id => {
        const e = document.getElementById(id); return e ? txt(e) : '';
      }).join(' ').trim();
      if (s) return s;
    }
    if (isInput) {
      if (el.id) {
        try {
          const la = document.querySelector('label[for="' + CSS.escape(el.id) + '"]');
          if (la) { const s = txt(la); if (s) return s; }
        } catch (e) {}
      }
      const w = el.closest('label');
      if (w) { const s = txt(w); if (s) return s; }
      let p = (el.getAttribute('placeholder') || el.getAttribute('aria-placeholder')
               || el.getAttribute('data-placeholder') || '').trim();
      if (p) return p;
      if (el.value) { const v = String(el.value).trim(); if (v) return v; }
      return (el.getAttribute('name') || el.getAttribute('title') || '').trim();
    }
    let s = txt(el); if (s) return s;
    if (el.value) { const v = String(el.value).trim(); if (v) return v; }
    let p = (el.getAttribute('placeholder') || '').trim(); if (p) return p;
    return (el.getAttribute('title') || el.getAttribute('name')
            || el.getAttribute('alt') || '').trim();
  };

  const sel = [
    'a[href]', 'button', 'input', 'textarea', 'select',
    '[role=button]', '[role=link]', '[role=textbox]', '[role=checkbox]',
    '[role=radio]', '[role=tab]', '[role=menuitem]', '[role=combobox]',
    '[role=searchbox]', '[role=switch]', '[contenteditable=true]',
    '[contenteditable=""]', '[onclick]'
  ].join(',');
  const seen = new Set();
  const out = [];
  let id = 0;
  for (const el of document.querySelectorAll(sel)) {
    if (seen.has(el)) continue;
    seen.add(el);
    const rect = el.getBoundingClientRect();
    if (rect.width < 2 || rect.height < 2) continue;
    const st = window.getComputedStyle(el);
    if (st.visibility === 'hidden' || st.display === 'none' || +st.opacity === 0) continue;
    if (el.disabled) continue;

    const tag = el.tagName.toLowerCase();
    const type = (el.getAttribute('type') || '').toLowerCase();
    const role = el.getAttribute('role') || '';
    const editable = el.isContentEditable;
    const isInput = tag === 'input' || tag === 'textarea' || tag === 'select'
                    || editable || role === 'textbox' || role === 'searchbox'
                    || role === 'combobox';

    let label = nameOf(el, isInput).replace(/\s+/g, ' ').slice(0, 90);

    // Skip noise: non-input elements with no readable label.
    if (!label && !isInput) continue;

    let kind;
    if (isInput) kind = (type && tag === 'input') ? ('input:' + type) : (editable ? 'editor' : tag);
    else kind = role || tag;

    el.setAttribute('data-jarvis-id', String(id));
    out.push({ id, kind, label, input: isInput });
    id++;
  }
  return out;
}
"""


def _index_interactive(page):
    """Tag every visible interactive element with data-jarvis-id and return a
    formatted numbered list plus the raw element data."""
    try:
        items = page.evaluate(_INDEX_JS)
    except Exception:
        items = []
    inputs, clicks = [], []
    for it in items:
        label = it.get("label", "") or "(no label)"
        line = f"  [{it['id']}] {it['kind']}: \"{label}\""
        (inputs if it.get("input") else clicks).append(line)
    parts = []
    if inputs:
        parts.append("FIELDS you can type into (browser_type element=number):\n" + "\n".join(inputs))
    if clicks:
        parts.append("BUTTONS/LINKS you can click (browser_click text=number):\n" + "\n".join(clicks))
    return "\n\n".join(parts), items


def _as_index(value):
    """Return an int element-id if *value* is a bare number (e.g. '5', '[5]'),
    else None. Lets the model address elements by their snapshot number."""
    s = str(value).strip().lstrip("[").rstrip("]").strip()
    if s.isdigit():
        return int(s)
    return None


def _settle(page, idle_ms: int = 4000, pause: float = 0.7) -> None:
    """Wait for a page to stop changing after an action. Clicks that open a
    modal/composer or load content need a beat before the DOM reflects the new
    state — snapshotting too early is why an action's result looks 'unchanged'.
    Generic: no per-site logic, just wait for network to go idle then pause for
    any open/close animation to finish."""
    try:
        page.wait_for_load_state("networkidle", timeout=idle_ms)
    except Exception:
        pass  # not all actions trigger network — the pause below still helps
    time.sleep(pause)


def _fill_locator(page, loc, content) -> bool:
    """Type *content* into a located element, handling both plain inputs
    (fill) and rich-text / contenteditable editors (click + keyboard type)."""
    try:
        loc.wait_for(state="visible", timeout=4000)
    except Exception:
        return False
    try:
        loc.scroll_into_view_if_needed(timeout=2000)
    except Exception:
        pass
    # Plain inputs/textarea/select accept fill() directly.
    try:
        tag = loc.evaluate("el => el.tagName.toLowerCase()")
    except Exception:
        tag = ""
    if tag in ("input", "textarea", "select"):
        try:
            loc.fill(content, timeout=3000)
            return True
        except Exception:
            pass
    # contenteditable / role=textbox — click to focus, then type.
    try:
        loc.click(timeout=3000)
        time.sleep(0.3)
        try:
            loc.fill("", timeout=1500)   # clear if it's fillable
        except Exception:
            pass
        page.keyboard.type(str(content), delay=12)
        return True
    except Exception:
        return False


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  TOOL DEFINITIONS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
TOOLS = [
    # ── quick URL open (no interaction needed) ────────────────────────────────
    {
        "type": "function",
        "function": {
            "name": "open_url",
            "description": (
                "Open a URL in the system default browser. Use this ONLY when you just "
                "want to show the user a page — no further interaction needed. "
                "For tasks that require typing, clicking, or reading the page, "
                "use browser_navigate + browser_* tools instead."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {"type": "string", "description": "Full URL with https://"}
                },
                "required": ["url"],
            },
        },
    },
    # ── OS / shell ────────────────────────────────────────────────────────────
    {
        "type": "function",
        "function": {
            "name": "run_shell_command",
            "description": (
                "Run a Windows cmd or PowerShell command. Use for opening native apps, "
                "file operations, system tasks. Do NOT pass raw URLs here — use open_url. "
                "Examples: 'start notepad.exe', 'start calc.exe', 'start spotify:', "
                "'start microsoft.windows.camera:', 'start ms-clock:Stopwatch', "
                "'shutdown /s /t 1', 'shutdown /r /t 1'."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "command":    {"type": "string",  "description": "Command to run"},
                    "powershell": {"type": "boolean", "description": "True = PowerShell, False = cmd"},
                },
                "required": ["command"],
            },
        },
    },
    # ── real-time web search ──────────────────────────────────────────────────
    {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": (
                "Search the web with DuckDuckGo and return current results. "
                "Use for: news, weather, prices, current events, anything that needs "
                "up-to-date information the model may not know."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query":       {"type": "string",  "description": "Search query"},
                    "max_results": {"type": "integer", "description": "Number of results (default 5)"},
                },
                "required": ["query"],
            },
        },
    },
    # ── browser automation ────────────────────────────────────────────────────
    {
        "type": "function",
        "function": {
            "name": "browser_navigate",
            "description": (
                "Navigate the Playwright browser to a URL. Use this (not open_url) "
                "when you need to interact with the page afterwards. "
                "Logins persist between sessions."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {"type": "string", "description": "Full URL with https://"}
                },
                "required": ["url"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "browser_snapshot",
            "description": (
                "Read the current browser page. Returns a NUMBERED list of every "
                "interactive element (buttons, links, inputs, editors) plus a text preview. "
                "Each element has a [number] — use that number with browser_click and "
                "browser_type. Call this after navigating or after any action that changes the page."
            ),
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "browser_click",
            "description": (
                "Click an element on the current page. STRONGLY PREFER passing the "
                "element's [number] from the latest browser_snapshot (e.g. '5'). "
                "You may pass visible text as a fallback, but numbers are far more reliable."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "text": {"type": "string", "description": "The element [number] from the snapshot (preferred), or visible text"}
                },
                "required": ["text"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "browser_type",
            "description": (
                "Type text into an input field or editor on the current page. "
                "STRONGLY PREFER passing the field's [number] from the latest "
                "browser_snapshot as 'element'. Set submit=true to press Enter after "
                "typing (useful for search boxes)."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "element": {"type": "string", "description": "The field's [number] from the snapshot (preferred), or its placeholder/label"},
                    "content": {"type": "string", "description": "Text to type"},
                    "submit":  {"type": "boolean", "description": "Press Enter after typing (default false)"},
                },
                "required": ["element", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "browser_key",
            "description": "Press a keyboard key in the browser. E.g. 'Enter', 'Tab', 'Escape', 'Control+a'.",
            "parameters": {
                "type": "object",
                "properties": {
                    "key": {"type": "string", "description": "Key name (Playwright key format)"}
                },
                "required": ["key"],
            },
        },
    },
    # ── native Windows app control ────────────────────────────────────────────
    {
        "type": "function",
        "function": {
            "name": "windows_control",
            "description": (
                "Control a native Windows application using the Windows Accessibility API. "
                "Actions: 'focus' (bring window to front), 'click' (click a button by text), "
                "'type' (type text into focused app or specific element), 'get_text' (read window content). "
                "Use for apps like Notepad, Calculator, Spotify desktop, File Explorer, etc."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "app_name": {"type": "string", "description": "Partial window title to match"},
                    "action":   {"type": "string", "description": "focus | click | type | get_text"},
                    "element":  {"type": "string", "description": "Button/element text (for click/type)"},
                    "text":     {"type": "string", "description": "Text to type (for type action)"},
                },
                "required": ["app_name", "action"],
            },
        },
    },
    # ── knowledge / utility ───────────────────────────────────────────────────
    {
        "type": "function",
        "function": {
            "name": "wikipedia_search",
            "description": "Search Wikipedia and return a 3-sentence summary. Use for facts, people, places, concepts.",
            "parameters": {
                "type": "object",
                "properties": {
                    "topic": {"type": "string", "description": "Topic to look up"}
                },
                "required": ["topic"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_current_time",
            "description": "Get the current local time.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_current_date",
            "description": "Get today's date.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "translate_text",
            "description": "Translate text to another language.",
            "parameters": {
                "type": "object",
                "properties": {
                    "text":            {"type": "string", "description": "Text to translate"},
                    "target_language": {"type": "string", "description": "Target language (name or ISO code)"},
                },
                "required": ["text", "target_language"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "shorten_url",
            "description": "Shorten a long URL via TinyURL.",
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {"type": "string", "description": "URL to shorten"}
                },
                "required": ["url"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "open_game",
            "description": (
                "ONLY use this tool when the user explicitly asks to play or open one of these specific games: "
                "breakout, car arcade, obstacle course, space invaders, football champs, path finder, tic tac toe. "
                "Do NOT use this for any other task — not for posting, searching, browsing, or anything else."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "game_name": {
                        "type": "string",
                        "description": "One of: breakout, car arcade, obstacle course, space invaders, football champs, path finder, tic tac toe, random",
                    }
                },
                "required": ["game_name"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "look_at_page",
            "description": (
                "SEE the current browser page with real computer vision (you have eyes). "
                "Use this to VISUALLY VERIFY the result of an action when the element list is "
                "not enough — e.g. after submitting, confirm a post actually published; read an "
                "error dialog, a CAPTCHA prompt, a chart, or any visual content the DOM text "
                "doesn't capture. Ask a specific question about what is on screen."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "question": {"type": "string",
                        "description": "What you want to know about the page, e.g. 'Did my post publish successfully? What does the confirmation say?'"}
                },
                "required": ["question"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "look_at_screen",
            "description": (
                "SEE the user's entire desktop screen with real computer vision. Use when the "
                "user asks about what is on their screen, to read an on-screen error/message, "
                "describe an image or app, or understand visual context outside the browser."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "question": {"type": "string",
                        "description": "What to look for or answer about the screen."}
                },
                "required": ["question"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "remember",
            "description": (
                "Save a durable fact about the user to long-term memory so you recall it in "
                "future conversations — preferences, names, accounts, recurring tasks, how they "
                "like things done. Use ONLY for lasting facts, never for one-off chit-chat. "
                "Relevant memories are surfaced to you automatically before each command."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "fact": {"type": "string",
                        "description": "The fact to remember, written as a standalone sentence, e.g. 'The user prefers concise answers and lives in Vellore.'"}
                },
                "required": ["fact"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "recall",
            "description": (
                "Search your long-term memory for things you know about the user that relate to "
                "a topic. Use when you need older context that wasn't already surfaced to you."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "What to look up in memory."}
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "click_by_vision",
            "description": (
                "Click something on the current browser page by DESCRIBING it visually, when "
                "it is NOT in the numbered element list (e.g. an image, a canvas control, an "
                "icon with no label, a map pin). Jarvis looks at the page and clicks where the "
                "thing is. Prefer browser_click with a number when the element IS listed."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "description": {"type": "string",
                        "description": "What to click, described visually, e.g. 'the blue circular play button in the centre of the video'."}
                },
                "required": ["description"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_weather",
            "description": "Get current weather for a place (free, no key). Use for weather questions.",
            "parameters": {
                "type": "object",
                "properties": {
                    "location": {"type": "string", "description": "City or place name, e.g. 'Vellore' or 'Tokyo'."}
                },
                "required": ["location"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "set_reminder",
            "description": (
                "Set a reminder/timer. Jarvis will speak the reminder out loud when it's due. "
                "Convert the user's phrasing into delay_seconds (e.g. 'in 10 minutes' → 600, "
                "'in 2 hours' → 7200)."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "text": {"type": "string", "description": "What to remind the user about."},
                    "delay_seconds": {"type": "integer", "description": "Seconds from now until the reminder fires."}
                },
                "required": ["text", "delay_seconds"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_clipboard",
            "description": (
                "Read whatever text the user currently has copied to their clipboard. Use when "
                "they say things like 'summarise what I copied', 'translate this', 'what does "
                "this mean' referring to copied text."
            ),
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "ask_file",
            "description": (
                "Read a local document (.txt .md .csv .pdf .docx or code) and answer a question "
                "about its contents. Use when the user asks about a file by path."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Full path to the file."},
                    "question": {"type": "string", "description": "What the user wants to know about it."}
                },
                "required": ["path", "question"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "calculate",
            "description": "Evaluate a math expression (arithmetic, powers, sqrt, sin/cos/log, etc.). Use for any calculation instead of doing mental math.",
            "parameters": {
                "type": "object",
                "properties": {
                    "expression": {"type": "string", "description": "e.g. '(45*1.18) + sqrt(144)' or '2**10'."}
                },
                "required": ["expression"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_news",
            "description": "Get current news headlines (free). Optionally about a topic. Use for 'what's the news', 'latest on X'.",
            "parameters": {
                "type": "object",
                "properties": {
                    "topic": {"type": "string", "description": "Optional topic, e.g. 'AI' or 'cricket'. Omit for top headlines."}
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "daily_briefing",
            "description": "Give a spoken briefing: greeting, date/time, weather, top headlines and pending reminders. Use for 'brief me' / 'what's my morning look like'.",
            "parameters": {
                "type": "object",
                "properties": {
                    "location": {"type": "string", "description": "City for the weather part (use what you know about the user if available)."},
                    "topic": {"type": "string", "description": "Optional news topic focus."}
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "set_plan",
            "description": (
                "At the START of a multi-step task, declare your plan as a list of short steps. "
                "It shows on the user's HUD as a live checklist. Then call complete_step after "
                "finishing each step. Use for tasks with 3+ steps."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "steps": {"type": "array", "items": {"type": "string"},
                              "description": "Ordered short step descriptions (max 8)."}
                },
                "required": ["steps"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "complete_step",
            "description": "Mark the plan step at the given index (0-based) as complete. Call right after finishing that step.",
            "parameters": {
                "type": "object",
                "properties": {
                    "index": {"type": "integer", "description": "0-based index of the step just completed."}
                },
                "required": ["index"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "browser_scroll",
            "description": (
                "Scroll the current page when what you need is below/above the fold. Returns a "
                "fresh snapshot of the now-visible elements. Use if a snapshot didn't show the "
                "element you expected."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "direction": {"type": "string", "enum": ["down", "up", "top", "bottom"],
                                  "description": "Which way to scroll."}
                },
                "required": ["direction"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "browser_tab",
            "description": "Manage browser tabs: open a new tab (optionally at a URL), list open tabs, or switch to a tab by index.",
            "parameters": {
                "type": "object",
                "properties": {
                    "action": {"type": "string", "enum": ["new", "list", "switch", "close"]},
                    "url": {"type": "string", "description": "URL for action=new (optional)."},
                    "index": {"type": "integer", "description": "Tab index for switch/close."}
                },
                "required": ["action"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "write_file",
            "description": "Create/overwrite a local file with text content (.txt .md .csv code, or .docx). Use to save notes, drafts, code or documents the user asks you to write.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Full destination path."},
                    "content": {"type": "string", "description": "The full file contents."}
                },
                "required": ["path", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "convert_currency",
            "description": "Convert an amount between currencies using live free exchange rates.",
            "parameters": {
                "type": "object",
                "properties": {
                    "amount": {"type": "number"},
                    "from_currency": {"type": "string", "description": "3-letter code, e.g. USD."},
                    "to_currency": {"type": "string", "description": "3-letter code, e.g. INR."}
                },
                "required": ["amount", "from_currency", "to_currency"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "media_control",
            "description": "Control system media playback (Spotify, YouTube, any player) via media keys.",
            "parameters": {
                "type": "object",
                "properties": {
                    "action": {"type": "string",
                               "enum": ["play", "pause", "next", "previous", "stop", "volup", "voldown", "mute"]}
                },
                "required": ["action"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_url",
            "description": "Fetch a specific web page/article by URL and summarise it or answer a question about it. Use when the user gives a link or says 'summarise this article'.",
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {"type": "string"},
                    "question": {"type": "string", "description": "Optional specific question; omit to summarise."}
                },
                "required": ["url"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_crypto_price",
            "description": "Get the live price of a cryptocurrency (free).",
            "parameters": {
                "type": "object",
                "properties": {
                    "coin": {"type": "string", "description": "e.g. 'bitcoin', 'eth', 'solana'."},
                    "currency": {"type": "string", "description": "fiat code, default usd."}
                },
                "required": ["coin"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "define_word",
            "description": "Get the dictionary definition of a word (free).",
            "parameters": {
                "type": "object",
                "properties": {"word": {"type": "string"}},
                "required": ["word"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "convert_units",
            "description": "Convert between units of length, mass, volume, speed, data, or temperature.",
            "parameters": {
                "type": "object",
                "properties": {
                    "value": {"type": "number"},
                    "from_unit": {"type": "string", "description": "e.g. 'km', 'lb', 'celsius'."},
                    "to_unit": {"type": "string", "description": "e.g. 'miles', 'kg', 'fahrenheit'."}
                },
                "required": ["value", "from_unit", "to_unit"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "deep_research",
            "description": (
                "Thoroughly research a topic across MULTIPLE web sources and give a synthesised, "
                "cited answer. Use for open questions needing real research ('research X', "
                "'compare Y vs Z', 'what's the consensus on…'), not simple lookups."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "topic": {"type": "string"},
                    "depth": {"type": "integer", "description": "How many sources to read (1-5, default 3)."}
                },
                "required": ["topic"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "watch_screen",
            "description": (
                "Turn proactive screen-watching ON or OFF. When on, Jarvis periodically glances "
                "at the screen and offers help if you seem stuck. Use when the user asks Jarvis "
                "to 'watch my screen' / 'keep an eye out' / 'stop watching'."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "enabled": {"type": "boolean"},
                    "interval_seconds": {"type": "integer", "description": "Seconds between glances (default 90, min 30)."}
                },
                "required": ["enabled"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "forget",
            "description": "Delete a stored long-term memory that matches a description. Use when the user says 'forget that' / 'forget what I said about X'.",
            "parameters": {
                "type": "object",
                "properties": {"query": {"type": "string", "description": "Which memory to remove."}},
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "computer_control",
            "description": (
                "Control ANY Windows application by sight (not just the browser) — click "
                "buttons/menus/icons or type into fields in Notepad, Settings, Spotify, VS Code, "
                "games, anything on screen. Jarvis looks at the foreground window, finds the "
                "control you describe, and acts on it. Use for desktop apps the browser tools "
                "can't reach."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "description": {"type": "string", "description": "The on-screen control to act on, described visually."},
                    "action": {"type": "string", "enum": ["click", "double_click", "right_click", "type"]},
                    "text": {"type": "string", "description": "Text to type (for action=type)."},
                    "window": {"type": "string", "description": "Optional: title (or part) of the window to focus first."}
                },
                "required": ["description", "action"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "teach_skill",
            "description": (
                "Create a brand-new permanent skill (tool) for yourself by writing Python. Use "
                "when the user asks you to 'learn to…' something repeatable you don't already have "
                "a tool for. The code MUST define `def run(args): -> str` where args is a dict with "
                "an 'input' string. Keep it self-contained (you may import stdlib + requests)."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "name": {"type": "string", "description": "short snake_case skill name."},
                    "description": {"type": "string", "description": "what the skill does + what to pass as input."},
                    "code": {"type": "string", "description": "Python defining def run(args): ... returning a string."}
                },
                "required": ["name", "description", "code"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "mission",
            "description": (
                "Enter AUTONOMOUS mission mode for a complex multi-step goal: Jarvis plans, "
                "executes, VERIFIES its own work with vision, and keeps going until the goal is "
                "truly done — no further input needed. Use for big 'just get this done' requests."
            ),
            "parameters": {
                "type": "object",
                "properties": {"goal": {"type": "string", "description": "The end goal to achieve autonomously."}},
                "required": ["goal"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "look_through_webcam",
            "description": (
                "Look at the user through their WEBCAM (which faces them) and answer a question "
                "about what it sees — what they're holding up, how they look, what they're doing, "
                "who/what is in view. Use for 'what do you see', 'what am I holding', 'how do I "
                "look', 'is anyone behind me'."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "question": {"type": "string", "description": "What to look for / answer about the webcam view."}
                },
                "required": ["question"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "check_posture",
            "description": "Use the side-facing webcam to assess the user's sitting posture and give a quick tip if they're slouching.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "presence_mode",
            "description": "Turn the webcam presence-awareness ON or OFF (it greets you when you arrive, notices when you leave, and shows a live view on the HUD). Turning it off releases the camera.",
            "parameters": {
                "type": "object",
                "properties": {"enabled": {"type": "boolean"}},
                "required": ["enabled"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "camera_recall",
            "description": (
                "Review the last ~couple minutes of webcam frames (kept locally in memory) to "
                "answer questions about what happened — 'did anyone come to my desk while I was "
                "away?', 'was I on my phone?', 'what was I doing earlier?'."
            ),
            "parameters": {
                "type": "object",
                "properties": {"question": {"type": "string"}},
                "required": ["question"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "focus_report",
            "description": "Report how long the user has been at their desk and how often they stepped away this session.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "take_photo",
            "description": "Snap a photo from the webcam and save it. Use for 'take a photo / picture of me'.",
            "parameters": {
                "type": "object",
                "properties": {"path": {"type": "string", "description": "Optional save path; defaults to Pictures."}},
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "scan_qr",
            "description": "Read a QR code the user holds up to the webcam (decodes locally). If it's a URL you can then open it.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "schedule_task",
            "description": (
                "Schedule a command to run LATER, automatically, through your full tools — "
                "either once after a delay, or every day at a set time. Jarvis runs it on its "
                "own and notifies the user (voice + phone). Use for 'every morning at 8…', "
                "'in 2 hours, …', 'remind+do X later'."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "command": {"type": "string", "description": "The natural-language command to run when it fires."},
                    "delay_seconds": {"type": "integer", "description": "For a one-shot: seconds from now."},
                    "at_time": {"type": "string", "description": "Clock time 'HH:MM' (24h) for a timed/daily task."},
                    "daily": {"type": "boolean", "description": "true = repeat every day at at_time."}
                },
                "required": ["command"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_scheduled",
            "description": "List the user's scheduled/automated tasks.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "cancel_scheduled",
            "description": "Cancel a scheduled task by its number (from list_scheduled).",
            "parameters": {
                "type": "object",
                "properties": {"index": {"type": "integer", "description": "0-based index from list_scheduled."}},
                "required": ["index"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "notify_phone",
            "description": "Push a message to the user's phone via Telegram (use to proactively tell them something when they may be away).",
            "parameters": {
                "type": "object",
                "properties": {"message": {"type": "string"}},
                "required": ["message"],
            },
        },
    },
] + [
    # ── EXTENDED TOOLBOX (compact schemas) ──────────────────────────────────
    {"type": "function", "function": {"name": n, "description": d, "parameters": {
        "type": "object",
        "properties": {k: ({"type": "integer", "description": k} if k in
                           ("count", "length", "sides", "days", "top", "lo", "hi", "level")
                           else {"type": "boolean", "description": k} if k in ("symbols",)
                           else {"type": "string", "description": k}) for k in props},
        "required": req}}}
    for n, d, props, req in [
        ("list_files", "List the files/folders in a directory.", ["path"], []),
        ("search_files", "Search for files by name under a folder (defaults to home).", ["query", "path"], ["query"]),
        ("open_folder", "Open a folder in Windows Explorer.", ["path"], ["path"]),
        ("create_folder", "Create a new folder.", ["path"], ["path"]),
        ("move_path", "Move a file or folder to a new location.", ["src", "dst"], ["src", "dst"]),
        ("copy_path", "Copy a file or folder.", ["src", "dst"], ["src", "dst"]),
        ("delete_path", "Delete a file/folder (safely moved to a trash folder, reversible).", ["path"], ["path"]),
        ("rename_path", "Rename a file or folder.", ["path", "new_name"], ["path", "new_name"]),
        ("zip_path", "Zip a file or folder.", ["path"], ["path"]),
        ("unzip_file", "Extract a .zip archive.", ["path", "dest"], ["path"]),
        ("disk_usage", "Report free/total disk space.", ["path"], []),
        ("file_info", "Size and modified-date of a file/folder.", ["path"], ["path"]),
        ("system_info", "OS, CPU, RAM and GPU summary.", [], []),
        ("list_processes", "Top running processes by memory.", ["top"], []),
        ("kill_process", "Close/terminate processes matching a name.", ["name"], ["name"]),
        ("battery_status", "Battery percentage and charging state.", [], []),
        ("screenshot_save", "Save a screenshot of the desktop to a file.", ["path"], []),
        ("set_clipboard", "Put text onto the clipboard.", ["text"], ["text"]),
        ("ip_info", "Your public IP address and rough location.", [], []),
        ("summarize_text", "Summarise a block of text concisely.", ["text"], ["text"]),
        ("rewrite_text", "Rewrite text in a given style/tone (style optional).", ["text", "style"], ["text"]),
        ("fix_grammar", "Fix spelling and grammar in text.", ["text"], ["text"]),
        ("generate_password", "Generate a strong random password.", ["length", "symbols"], []),
        ("hash_text", "Hash text (md5/sha1/sha256).", ["text", "algo"], ["text"]),
        ("base64_tool", "Base64 encode or decode text (mode=encode|decode).", ["text", "mode"], ["text"]),
        ("format_json", "Pretty-print / validate JSON.", ["text"], ["text"]),
        ("count_words", "Word/character/line counts for text.", ["text"], ["text"]),
        ("qr_generate", "Make a QR-code image for a link or text.", ["data", "path"], ["data"]),
        ("describe_image", "Describe / answer about a local image file (vision).", ["path", "question"], ["path"]),
        ("get_forecast", "Multi-day weather forecast for a place.", ["location", "days"], ["location"]),
        ("air_quality", "Air-quality index for a place.", ["location"], ["location"]),
        ("sunrise_sunset", "Sunrise and sunset times for a place.", ["location"], ["location"]),
        ("hacker_news", "Top Hacker News stories.", ["count"], []),
        ("github_repo", "Stats for a GitHub repo (owner/name).", ["repo"], ["repo"]),
        ("synonyms", "Synonyms for a word.", ["word"], ["word"]),
        ("random_fact", "A random interesting fact.", [], []),
        ("tell_joke", "Tell a random joke.", [], []),
        ("this_day", "Notable events on this day in history.", [], []),
        ("stock_price", "Current price of a stock ticker (e.g. AAPL).", ["symbol"], ["symbol"]),
        ("time_in", "Current local time in a city/timezone.", ["city"], ["city"]),
        ("roll_dice", "Roll dice (sides, count).", ["sides", "count"], []),
        ("flip_coin", "Flip a coin.", [], []),
        ("random_number", "Random number between lo and hi.", ["lo", "hi"], []),
        ("days_until", "Days until/since a date (YYYY-MM-DD).", ["date"], ["date"]),
        ("expand_url", "Resolve a shortened URL to its real destination.", ["url"], ["url"]),
    ]
]


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  TOOL EXECUTOR
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
GAME_FILES = {
    "breakout":        "breakout arcade gaming.py",
    "car arcade":      "car arcade gaming.py",
    "obstacle":        "complete obby game! - ursina.py",
    "obstacle course": "complete obby game! - ursina.py",
    "space invaders":  "space-invaders final game.py",
    "football":        "football champs.py",
    "football champs": "football champs.py",
    "path finder":     "path finder.py",
    "tic tac toe":     "tic tac toe2options.py",
}


def execute_tool(name: str, args: dict) -> str:
    """Thin wrapper so NO tool can ever crash the agent loop — any unexpected
    exception (bad numeric arg, permission error, etc.) becomes a friendly string."""
    print(f"⚙  Tool → {name}({args})")
    try:
        return _execute_tool_impl(name, args or {})
    except Exception as e:
        return f"Tool '{name}' hit an error: {e}"


def _execute_tool_impl(name: str, args: dict) -> str:
    # ── open_url ──────────────────────────────────────────────────────────────
    if name == "open_url":
        url = args.get("url", "").strip()
        if not url.startswith("http"):
            url = "https://" + url
        webbrowser.open(url)
        return f"Opened {url}"

    # ── run_shell_command ─────────────────────────────────────────────────────
    elif name == "run_shell_command":
        cmd = args.get("command", "")
        use_ps = args.get("powershell", False)
        if cmd.startswith("http://") or cmd.startswith("https://"):
            webbrowser.open(cmd)
            return f"Opened {cmd}"
        try:
            if use_ps:
                proc = subprocess.run(
                    ["powershell", "-Command", cmd],
                    capture_output=True, text=True, timeout=15,
                )
            else:
                proc = subprocess.run(
                    cmd, shell=True, capture_output=True, text=True, timeout=15,
                )
            out = (proc.stdout or proc.stderr or "Done.").strip()
            return out[:500]
        except subprocess.TimeoutExpired:
            return "Started (running in background)."
        except Exception as e:
            return f"Error: {e}"

    # ── web_search ────────────────────────────────────────────────────────────
    elif name == "web_search":
        query = args.get("query", "")
        max_r = int(args.get("max_results", 5))
        try:
            from ddgs import DDGS
            results = []
            with DDGS() as ddgs:
                for r in ddgs.text(query, max_results=max_r):
                    results.append(
                        f"• {r['title']}\n  {r['href']}\n  {r['body'][:200]}"
                    )
            return "\n\n".join(results) if results else "No results found."
        except Exception as e:
            return f"Search error: {e}"

    # ── browser_navigate ──────────────────────────────────────────────────────
    elif name == "browser_navigate":
        url = args.get("url", "")
        if not url.startswith("http"):
            url = "https://" + url
        try:
            page = _get_page()
            page.goto(url, wait_until="domcontentloaded", timeout=25000)
            # Give JS-heavy pages (LinkedIn, GitHub, etc.) time to render
            try:
                page.wait_for_load_state("networkidle", timeout=5000)
            except Exception:
                pass  # timeout is fine — page is usable, just still loading extras
            time.sleep(0.5)
            snap = execute_tool("browser_snapshot", {})
            return f"Navigated to: {page.title()} ({page.url})\n\n{snap}"
        except Exception as e:
            return f"Navigation error: {e}"

    # ── browser_snapshot ──────────────────────────────────────────────────────
    elif name == "browser_snapshot":
        try:
            page = _get_page()
            listing, items = _index_interactive(page)
            title   = page.title()
            url_now = page.url
            try:
                preview = page.inner_text("body")[:700].replace("\n", " ").strip()
            except Exception:
                preview = ""

            result  = f"Page: {title}\nURL:  {url_now}\n\n"
            if items:
                result += listing
            else:
                result += "No interactive elements detected (page may still be loading)."
            result += f"\n\nPage text preview:\n{preview}"
            return result
        except Exception as e:
            return f"Snapshot error: {e}"

    # ── browser_click ─────────────────────────────────────────────────────────
    elif name == "browser_click":
        text = str(args.get("text", "")).strip()
        try:
            page = _get_page()

            # Preferred path: the model gave us an element NUMBER from the snapshot.
            idx = _as_index(text)
            if idx is not None:
                loc = page.locator(f"[data-jarvis-id='{idx}']")
                if loc.count() == 0:
                    return (
                        f"No element [{idx}] on this page (numbers change when the page "
                        f"updates). Call browser_snapshot to get fresh numbers."
                    )
                try:
                    loc.first.scroll_into_view_if_needed(timeout=2000)
                except Exception:
                    pass
                loc.first.click(timeout=5000)
                _settle(page)
                snap = execute_tool("browser_snapshot", {})
                return f"Clicked element [{idx}]. Page now:\n\n{snap}"

            # Fallback: text-based matching (kept for resilience).
            strategies = [
                lambda: page.get_by_role("button", name=text).first.click(timeout=3500),
                lambda: page.get_by_role("link",   name=text).first.click(timeout=3500),
                lambda: page.get_by_text(text, exact=False).first.click(timeout=3500),
                lambda: page.get_by_label(text).first.click(timeout=3500),
                lambda: page.get_by_placeholder(text).first.click(timeout=3500),
                lambda: page.locator(f"[aria-label*='{text}']").first.click(timeout=3500),
            ]
            for strategy in strategies:
                try:
                    strategy()
                    _settle(page)
                    snap = execute_tool("browser_snapshot", {})
                    return f"Clicked '{text}'. Page now:\n\n{snap}"
                except Exception:
                    continue
            # Self-heal: text matching failed → fall back to vision (Set-of-Marks).
            try:
                _hud_emit("toast", text="👁 click self-healing via vision…", level="info")
                vres = _vision_click(page, text)
                if vres.startswith("Saw and clicked"):
                    return vres
            except Exception:
                pass
            return (
                f"Could not find '{text}' to click, even with vision. Call browser_snapshot and "
                f"click using the element [number] instead."
            )
        except Exception as e:
            return f"Click error: {e}"

    # ── browser_type ──────────────────────────────────────────────────────────
    elif name == "browser_type":
        element = str(args.get("element", "")).strip()
        content = args.get("content", "")
        submit  = bool(args.get("submit", False))

        if not element:
            return (
                "ERROR: 'element' is empty. Call browser_snapshot, then pass the "
                "input field's [number] as 'element'."
            )

        try:
            page = _get_page()

            # Preferred path: the model gave us an element NUMBER from the snapshot.
            idx = _as_index(element)
            if idx is not None:
                base = page.locator(f"[data-jarvis-id='{idx}']")
                if base.count() == 0:
                    return (
                        f"No element [{idx}] on this page (numbers change when the page "
                        f"updates). Call browser_snapshot to get fresh numbers."
                    )
                loc = base.first
                ok = _fill_locator(page, loc, content)
                if not ok:
                    return f"Could not type into element [{idx}]."
                if submit:
                    page.keyboard.press("Enter")
                    time.sleep(0.6)
                _settle(page)
                snap = execute_tool("browser_snapshot", {})
                done = "Typed into element [{}].{}".format(idx, " Submitted." if submit else "")
                return f"{done} Page now:\n\n{snap}"

            # Fallback: text/label-based matching.
            fill_strategies = [
                lambda: page.get_by_placeholder(element).first,
                lambda: page.get_by_label(element).first,
                lambda: page.get_by_role("textbox",   name=element).first,
                lambda: page.get_by_role("searchbox", name=element).first,
            ]
            for make in fill_strategies:
                try:
                    loc = make()
                    if _fill_locator(page, loc, content):
                        if submit:
                            page.keyboard.press("Enter")
                            time.sleep(0.6)
                        _settle(page)
                        snap = execute_tool("browser_snapshot", {})
                        done = "Typed into '{}'.{}".format(element, " Submitted." if submit else "")
                        return f"{done} Page now:\n\n{snap}"
                except Exception:
                    continue
            return (
                f"Could not find field '{element}'. Call browser_snapshot and pass "
                f"the field's [number] as 'element'."
            )
        except Exception as e:
            return f"Type error: {e}"

    # ── browser_key ───────────────────────────────────────────────────────────
    elif name == "browser_key":
        key = args.get("key", "")
        try:
            page = _get_page()
            page.keyboard.press(key)
            time.sleep(0.3)
            return f"Pressed '{key}'"
        except Exception as e:
            return f"Key error: {e}"

    # ── windows_control ───────────────────────────────────────────────────────
    elif name == "windows_control":
        app_name = args.get("app_name", "")
        action   = args.get("action", "")
        element  = args.get("element", "")
        text     = args.get("text", "")

        # pywinauto/comtypes must run in a dedicated COM-initialised thread, or it
        # crashes with "Cannot change thread mode" because pygame already set the
        # main thread's COM apartment (same fix as _desktop_marks_data).
        out = {}

        def _work():
            try:
                try:
                    import comtypes
                    comtypes.CoInitialize()
                except Exception:
                    pass
                from pywinauto import Application, Desktop
                from pywinauto.keyboard import send_keys
                if action == "focus":
                    wins = Desktop(backend="uia").windows(title_re=f".*{app_name}.*")
                    if wins:
                        wins[0].set_focus()
                        out["r"] = f"Focused: {wins[0].window_text()}"
                    else:
                        out["r"] = f"No window found matching '{app_name}'"
                elif action == "click":
                    app = Application(backend="uia").connect(title_re=f".*{app_name}.*", timeout=5)
                    app.top_window().child_window(title=element, control_type="Button").click_input()
                    out["r"] = f"Clicked '{element}' in {app_name}"
                elif action == "type":
                    app = Application(backend="uia").connect(title_re=f".*{app_name}.*", timeout=5)
                    win = app.top_window(); win.set_focus()
                    if element:
                        win.child_window(title=element).set_edit_text(text)
                    else:
                        send_keys(text, with_spaces=True)
                    out["r"] = f"Typed in {app_name}"
                elif action == "get_text":
                    app = Application(backend="uia").connect(title_re=f".*{app_name}.*", timeout=5)
                    out["r"] = app.top_window().window_text()[:500]
                else:
                    out["r"] = f"Unknown action '{action}'"
            except Exception as e:
                out["r"] = f"Windows control error: {e}"

        t = threading.Thread(target=_work, daemon=True)
        t.start(); t.join(timeout=20)
        return out.get("r", "Windows control timed out.")

    # ── wikipedia_search ──────────────────────────────────────────────────────
    elif name == "wikipedia_search":
        topic = args.get("topic", "")
        try:
            import wikipedia
            summary = wikipedia.summary(topic, sentences=3)
            return summary
        except wikipedia.exceptions.DisambiguationError as e:
            return f"Multiple results for '{topic}'. Did you mean: {', '.join(e.options[:3])}?"
        except wikipedia.exceptions.PageError:
            return f"No Wikipedia page found for '{topic}'."
        except Exception as e:
            return f"Wikipedia error: {e}"

    # ── get_current_time ──────────────────────────────────────────────────────
    elif name == "get_current_time":
        return datetime.datetime.now().strftime("%I:%M %p")

    # ── get_current_date ──────────────────────────────────────────────────────
    elif name == "get_current_date":
        return date.today().strftime("%B %d, %Y")

    # ── translate_text ────────────────────────────────────────────────────────
    elif name == "translate_text":
        text   = args.get("text", "")
        target = args.get("target_language", "english")
        try:
            from deep_translator import GoogleTranslator
            return GoogleTranslator(source="auto", target=target).translate(text)
        except Exception as e:
            return f"Translation error: {e}"

    # ── shorten_url ───────────────────────────────────────────────────────────
    elif name == "shorten_url":
        url = args.get("url", "")
        try:
            import pyshorteners
            short = pyshorteners.Shortener().tinyurl.short(url)
            print(f"   Shortened URL: {short}")
            return short
        except Exception as e:
            return f"Could not shorten URL: {e}"

    # ── open_game ─────────────────────────────────────────────────────────────
    elif name == "open_game":
        game_name = args.get("game_name", "").lower()
        if "random" in game_name:
            game_file = _random.choice(list(GAME_FILES.values()))
        else:
            game_file = next(
                (v for k, v in GAME_FILES.items() if k in game_name), None
            )
        if not game_file:
            return f"Game '{game_name}' not found. Available: {', '.join(GAME_FILES)}"
        try:
            os.startfile(game_file)
            return f"Opened {game_file}"
        except FileNotFoundError:
            return f"Game file not found: {game_file}"

    # ── look_at_page ── (vision: see the live browser page) ─────────────────────
    elif name == "look_at_page":
        question = args.get("question", "What is on this page?")
        try:
            page = _get_page()
            png = page.screenshot(type="png")
            _hud_emit("toast", text="👁 Looking at the page…", level="info")
            return _gemini_vision(question, base64.b64encode(png).decode(), "image/png")
        except Exception as e:
            return f"Vision error (page): {e}"

    # ── look_at_screen ── (vision: see the whole desktop) ───────────────────────
    elif name == "look_at_screen":
        question = args.get("question", "What is on the screen?")
        try:
            b64, mime = _grab_screen_b64()
            _hud_emit("toast", text="👁 Looking at your screen…", level="info")
            return _gemini_vision(question, b64, mime)
        except Exception as e:
            return f"Vision error (screen): {e}"

    # ── remember ── (write to long-term semantic memory) ────────────────────────
    elif name == "remember":
        fact = (args.get("fact") or "").strip()
        if not fact:
            return "Nothing to remember (empty fact)."
        ok = _mem_add(fact)
        _hud_emit("toast", text="🧠 Remembered", level="ok")
        _hud_emit("memory", count=len(_MEM))
        return "Saved to long-term memory." if ok else "Couldn't reach the memory embedder; not saved."

    # ── recall ── (semantic search of long-term memory) ─────────────────────────
    elif name == "recall":
        query = (args.get("query") or "").strip()
        hits = _mem_search(query, k=5)
        if not hits:
            return "I don't have anything in memory about that yet."
        return "From memory:\n" + "\n".join(f"• {t}" for t, _ in hits)

    # ── click_by_vision ── (Set-of-Marks: model picks a marked element to click) ─
    elif name == "click_by_vision":
        desc = args.get("description", "")
        try:
            _hud_emit("toast", text="👁 locating by vision…", level="info")
            return _vision_click(_get_page(), desc)
        except Exception as e:
            return f"Vision-click error: {e}"

    # ── get_weather ─────────────────────────────────────────────────────────────
    elif name == "get_weather":
        return _get_weather(args.get("location", ""))

    # ── set_reminder ────────────────────────────────────────────────────────────
    elif name == "set_reminder":
        text = (args.get("text") or "").strip() or "your reminder"
        try:
            delay = int(args.get("delay_seconds", 0))
        except Exception:
            delay = 0
        if delay <= 0:
            return "I need a valid time for the reminder."
        return _set_reminder(text, delay)

    # ── read_clipboard ──────────────────────────────────────────────────────────
    elif name == "read_clipboard":
        clip = _read_clipboard()
        if not clip.strip():
            return "The clipboard is empty (or contains no text)."
        return f"Clipboard contents:\n{clip[:4000]}"

    # ── ask_file ──────────────────────────────────────────────────────────────--
    elif name == "ask_file":
        return _ask_file(args.get("path", ""), args.get("question", "Summarise this file."))

    # ── calculate ───────────────────────────────────────────────────────────────
    elif name == "calculate":
        return _calculate(args.get("expression", ""))

    # ── get_news ──────────────────────────────────────────────────────────────--
    elif name == "get_news":
        heads = _get_news(args.get("topic", ""), n=6)
        return ("Headlines:\n" + "\n".join(f"• {h}" for h in heads)) if heads else "Couldn't fetch news right now."

    # ── daily_briefing ──────────────────────────────────────────────────────────
    elif name == "daily_briefing":
        return _compose_briefing(args.get("location", ""), args.get("topic", ""))

    # ── set_plan / complete_step ────────────────────────────────────────────────
    elif name == "set_plan":
        return _set_plan(args.get("steps", []))
    elif name == "complete_step":
        return _complete_step(args.get("index", 0))

    # ── browser_scroll ──────────────────────────────────────────────────────────
    elif name == "browser_scroll":
        d = (args.get("direction") or "down").lower()
        try:
            page = _get_page()
            js = {"down": "window.scrollBy(0, window.innerHeight*0.85)",
                  "up": "window.scrollBy(0, -window.innerHeight*0.85)",
                  "top": "window.scrollTo(0, 0)",
                  "bottom": "window.scrollTo(0, document.body.scrollHeight)"}.get(d,
                  "window.scrollBy(0, window.innerHeight*0.85)")
            page.evaluate(js)
            _settle(page, idle_ms=1500, pause=0.4)
            snap = execute_tool("browser_snapshot", {})
            return f"Scrolled {d}. Page now:\n\n{snap}"
        except Exception as e:
            return f"Scroll error: {e}"

    # ── browser_tab ──────────────────────────────────────────────────────────---
    elif name == "browser_tab":
        global _pw_page
        action = (args.get("action") or "list").lower()
        try:
            page = _get_page()
            ctx = page.context
            if action == "new":
                np = ctx.new_page()
                _pw_page = np
                url = args.get("url", "")
                if url:
                    if not url.startswith("http"):
                        url = "https://" + url
                    np.goto(url, wait_until="domcontentloaded", timeout=25000)
                    _settle(np)
                return f"Opened a new tab.\n\n{execute_tool('browser_snapshot', {})}"
            pages = [p for p in ctx.pages if not p.is_closed()]
            if action == "list":
                return "Open tabs:\n" + "\n".join(
                    f"  [{i}] {p.title()[:50]}  ({p.url[:60]})" for i, p in enumerate(pages))
            idx = int(args.get("index", 0))
            if not (0 <= idx < len(pages)):
                return f"No tab [{idx}]. There are {len(pages)} tabs."
            if action == "switch":
                _pw_page = pages[idx]
                _pw_page.bring_to_front()
                return f"Switched to tab [{idx}].\n\n{execute_tool('browser_snapshot', {})}"
            if action == "close":
                pages[idx].close()
                _pw_page = None
                return f"Closed tab [{idx}]."
            return f"Unknown tab action '{action}'."
        except Exception as e:
            return f"Tab error: {e}"

    # ── write_file ──────────────────────────────────────────────────────────────
    elif name == "write_file":
        return _write_file(args.get("path", ""), args.get("content", ""))

    # ── convert_currency ─────────────────────────────────────────────────────────
    elif name == "convert_currency":
        try:
            amt = float(args.get("amount", 0))
        except Exception:
            amt = 0.0
        return _convert_currency(amt, args.get("from_currency", ""), args.get("to_currency", ""))

    # ── media_control ─────────────────────────────────────────────────────────--
    elif name == "media_control":
        return _media_control(args.get("action", ""))

    # ── read_url ─────────────────────────────────────────────────────────────--
    elif name == "read_url":
        return _read_url(args.get("url", ""), args.get("question", ""))

    # ── get_crypto_price ─────────────────────────────────────────────────────---
    elif name == "get_crypto_price":
        return _get_crypto_price(args.get("coin", ""), args.get("currency", "usd"))

    # ── define_word ──────────────────────────────────────────────────────────---
    elif name == "define_word":
        return _define_word(args.get("word", ""))

    # ── convert_units ────────────────────────────────────────────────────────---
    elif name == "convert_units":
        try:
            val = float(args.get("value", 0))
        except Exception:
            val = 0.0
        return _convert_units(val, args.get("from_unit", ""), args.get("to_unit", ""))

    # ── deep_research ────────────────────────────────────────────────────────---
    elif name == "deep_research":
        return _deep_research(args.get("topic", ""), args.get("depth", 3))

    # ── watch_screen ─────────────────────────────────────────────────────────---
    elif name == "watch_screen":
        _WATCH["on"] = bool(args.get("enabled", False))
        if args.get("interval_seconds"):
            try:
                _WATCH["interval"] = max(30, int(args["interval_seconds"]))
            except Exception:
                pass
        _hud_emit("toast", text=("👁 watching your screen" if _WATCH["on"] else "watch mode off"),
                  level="info")
        _hud_emit("watch", on=_WATCH["on"])
        return ("Watching your screen — I'll speak up if I can help."
                if _WATCH["on"] else "Stopped watching your screen.")

    # ── forget ───────────────────────────────────────────────────────────────---
    elif name == "forget":
        return _forget(args.get("query", ""))

    # ── computer_control (any desktop app by sight) ─────────────────────────────
    elif name == "computer_control":
        return _computer_control(args.get("description", ""), args.get("action", "click"),
                                 args.get("text", ""), args.get("window", ""))

    # ── teach_skill (write a new tool) ──────────────────────────────────────────
    elif name == "teach_skill":
        return _teach_skill(args.get("name", ""), args.get("description", ""), args.get("code", ""))

    # ── mission (autonomous mode) ───────────────────────────────────────────────
    elif name == "mission":
        return _mission(args.get("goal", ""))

    # ── webcam vision ────────────────────────────────────────────────────────---
    elif name == "look_through_webcam":
        _hud_emit("toast", text="📷 looking at you…", level="info")
        return _cam_vision(args.get("question", "What do you see?"))
    elif name == "check_posture":
        _hud_emit("toast", text="📷 checking posture…", level="info")
        return _cam_vision("This is a side view of the user at their desk. Assess their sitting "
                           "posture in one sentence and give a quick tip if they're slouching.")
    elif name == "presence_mode":
        _PRESENCE["on"] = bool(args.get("enabled", True))
        if not _PRESENCE["on"]:
            _cam_release()
        _hud_emit("presence", present=_PRESENCE.get("present", False), on=_PRESENCE["on"])
        return ("Presence awareness on — I'll keep an eye out." if _PRESENCE["on"]
                else "Presence awareness off; camera released.")
    elif name == "camera_recall":
        _hud_emit("toast", text="📷 reviewing recent frames…", level="info")
        return _camera_recall(args.get("question", "What happened recently?"))
    elif name == "focus_report":
        return _focus_report()
    elif name == "take_photo":
        path, err = _cam_photo(args.get("path", ""))
        return f"Saved a photo to {path}." if path else f"Couldn't take a photo: {err}"
    elif name == "scan_qr":
        _hud_emit("toast", text="📷 scanning QR…", level="info")
        data = _scan_qr()
        return f"QR code: {data}" if data else "I didn't see a QR code — hold it steady, facing the webcam."
    elif name == "schedule_task":
        return _schedule_task(args.get("command", ""), args.get("delay_seconds", 0),
                              args.get("at_time", ""), bool(args.get("daily", False)))
    elif name == "list_scheduled":
        with _SCHED_LOCK:
            if not _SCHEDULE:
                return "You have no scheduled tasks."
            out = []
            for i, it in enumerate(_SCHEDULE):
                when = (f"daily {it['time']}" if it.get("daily") else it.get("time")) if it.get("time") \
                       else "once soon"
                out.append(f"[{i}] {it.get('command','')[:50]} — {when}")
        return "Scheduled tasks:\n" + "\n".join(out)
    elif name == "cancel_scheduled":
        try:
            i = int(args.get("index", -1))
        except Exception:
            i = -1
        with _SCHED_LOCK:
            if 0 <= i < len(_SCHEDULE):
                gone = _SCHEDULE.pop(i); _sched_save()
                return f"Cancelled: {gone.get('command','')[:50]}"
        return "No task at that number."
    elif name == "notify_phone":
        m_txt = args.get("message", "")
        if not (JARVIS_TG_TOKEN and JARVIS_TG_CHAT):
            return "Phone isn't linked (set up the Telegram bot first)."
        _tg_send("🔔 " + m_txt)
        return "Sent to your phone."

    # ── EXTENDED TOOLBOX dispatch ────────────────────────────────────────────────
    elif name == "list_files":      return _list_files(args.get("path", ""))
    elif name == "search_files":    return _search_files(args.get("query", ""), args.get("path", ""))
    elif name == "open_folder":     return _open_folder(args.get("path", ""))
    elif name == "create_folder":   return _create_folder(args.get("path", ""))
    elif name == "move_path":       return _move_path(args.get("src", ""), args.get("dst", ""))
    elif name == "copy_path":       return _copy_path(args.get("src", ""), args.get("dst", ""))
    elif name == "delete_path":     return _delete_path(args.get("path", ""))
    elif name == "rename_path":     return _rename_path(args.get("path", ""), args.get("new_name", ""))
    elif name == "zip_path":        return _zip_path(args.get("path", ""))
    elif name == "unzip_file":      return _unzip_file(args.get("path", ""), args.get("dest", ""))
    elif name == "disk_usage":      return _disk_usage(args.get("path", ""))
    elif name == "file_info":       return _file_info(args.get("path", ""))
    elif name == "system_info":     return _system_info()
    elif name == "list_processes":  return _list_processes(int(args.get("top", 10) or 10))
    elif name == "kill_process":    return _kill_process(args.get("name", ""))
    elif name == "battery_status":  return _battery_status()
    elif name == "screenshot_save": return _screenshot_save(args.get("path", ""))
    elif name == "set_clipboard":   return _set_clipboard(args.get("text", ""))
    elif name == "ip_info":         return _ip_info()
    elif name == "summarize_text":  return _ai_text("Summarise this concisely:", args.get("text", ""))
    elif name == "rewrite_text":    return _ai_text(f"Rewrite this in a {args.get('style','clear, polished')} style:", args.get("text", ""))
    elif name == "fix_grammar":     return _ai_text("Fix the spelling and grammar; return only the corrected text:", args.get("text", ""))
    elif name == "generate_password": return _generate_password(int(args.get("length", 16) or 16), bool(args.get("symbols", True)))
    elif name == "hash_text":       return _hash_text(args.get("text", ""), args.get("algo", "sha256"))
    elif name == "base64_tool":     return _base64_tool(args.get("text", ""), args.get("mode", "encode"))
    elif name == "format_json":     return _format_json(args.get("text", ""))
    elif name == "count_words":     return _count_words(args.get("text", ""))
    elif name == "qr_generate":     return _qr_generate(args.get("data", ""), args.get("path", ""))
    elif name == "describe_image":  return _describe_image(args.get("path", ""), args.get("question", "Describe this image."))
    elif name == "get_forecast":    return _get_forecast(args.get("location", ""), int(args.get("days", 3) or 3))
    elif name == "air_quality":     return _air_quality(args.get("location", ""))
    elif name == "sunrise_sunset":  return _sunrise_sunset(args.get("location", ""))
    elif name == "hacker_news":     return _hacker_news(int(args.get("count", 5) or 5))
    elif name == "github_repo":     return _github_repo(args.get("repo", ""))
    elif name == "synonyms":        return _synonyms(args.get("word", ""))
    elif name == "random_fact":     return _random_fact()
    elif name == "tell_joke":       return _tell_joke()
    elif name == "this_day":        return _this_day()
    elif name == "stock_price":     return _stock_price(args.get("symbol", ""))
    elif name == "time_in":         return _time_in(args.get("city", ""))
    elif name == "roll_dice":       return _roll_dice(int(args.get("sides", 6) or 6), int(args.get("count", 1) or 1))
    elif name == "flip_coin":       return _flip_coin()
    elif name == "random_number":   return _random_number(int(args.get("lo", 1) or 1), int(args.get("hi", 100) or 100))
    elif name == "days_until":      return _days_until(args.get("date", ""))
    elif name == "expand_url":      return _expand_url(args.get("url", ""))

    # ── self-authored skills ────────────────────────────────────────────────────
    elif name in _SKILLS:
        try:
            return str(_SKILLS[name]["func"](args))
        except Exception as e:
            return f"Skill '{name}' errored: {e}"

    return f"Unknown tool: {name}"


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  SYSTEM PROMPT
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SYSTEM_PROMPT = """You are Jarvis, a personal AI voice assistant running on Windows 11.

━━ CORE RULES ━━
• Your responses are SPOKEN ALOUD — keep final replies SHORT (1-2 sentences).
• NEVER use markdown, bullet points, asterisks, or headers in your spoken replies.
• Always USE TOOLS to complete tasks — never pretend you did something without calling a tool.
• Never invent fake results. If a tool fails, say so honestly.

━━ TOOL SELECTION ━━
Task                                    → Tool
Open a site (no interaction needed)     → open_url
Post / create / fill / click on a site  → browser_navigate then browser_snapshot then browser_click/type
Read or interact with a webpage         → browser_navigate + browser_snapshot + browser_click/type
Real-time info (news/weather/prices)    → web_search
Open a native Windows app               → run_shell_command
Control an open Windows app             → windows_control
Facts / encyclopedia                    → wikipedia_search
See / read the user's screen            → look_at_screen
Visually verify a web action worked     → look_at_page
Click an unlabelled/visual element      → click_by_vision
Save a lasting fact about the user      → remember
Look up older context about the user    → recall
Weather for a place                     → get_weather
News headlines                          → get_news
A spoken briefing of the day            → daily_briefing
Any math / calculation                  → calculate
Set a reminder / timer                  → set_reminder
Use text the user has copied            → read_clipboard
Answer questions about a local file     → ask_file
Save/write a file or document           → write_file
Convert currency                        → convert_currency
Play/pause/next music or volume         → media_control
Reveal off-screen page content          → browser_scroll
Work across multiple tabs               → browser_tab
Summarise an article / read a link      → read_url
Crypto price                            → get_crypto_price
Define a word                           → define_word
Convert units (length/mass/temp/…)      → convert_units
Deep multi-source research              → deep_research
Watch the screen / offer proactive help → watch_screen
Forget a stored memory                  → forget
Control a non-browser desktop app       → computer_control
Learn a new repeatable ability          → teach_skill
Autonomously complete a big goal        → mission
See the user via their webcam           → look_through_webcam
Check the user's sitting posture        → check_posture
Recall recent webcam events             → camera_recall
How long at desk / focus stats          → focus_report
Take a photo with the webcam            → take_photo
Read a QR code shown to the webcam      → scan_qr
Do something later / every day          → schedule_task (list_scheduled, cancel_scheduled)
Push a message to the user's phone       → notify_phone

━━ YOU CAN SEE THE USER ━━
A webcam faces the user. look_through_webcam lets you SEE them — use it for "what
do you see / what am I holding / how do I look / is someone behind me". You also
sense their presence (you know when they arrive or step away). Be natural about it.

━━ SHOW YOUR PLAN ━━
For any task with 3 or more steps, FIRST call set_plan with the ordered steps — it
appears on the user's HUD as a live checklist. Call complete_step(index) right after
each step finishes so they can watch progress. Keep steps short.

━━ YOU HAVE EYES (vision) ━━
You can SEE. After an important web action whose success isn't obvious from the element
list — submitting a form, publishing a post, a payment/confirmation screen — call
look_at_page to VISUALLY CONFIRM it actually worked before telling the user it's done.
Use look_at_screen when the user asks about anything on their screen. This works on any
site or app; never assume success, verify it.

━━ YOU HAVE MEMORY ━━
Durable facts you should recall later (preferences, names, accounts, how the user likes
things done, recurring tasks) → call remember. Relevant memories are automatically given
to you before each command under "THINGS YOU REMEMBER"; use them naturally. Use recall
for older context that wasn't surfaced. Do NOT remember trivial one-off chit-chat.

━━ BROWSER WORKFLOW (works on ANY website) ━━
Every browser_snapshot returns a NUMBERED list of the page's interactive elements:
    [0] button: "Start a post"
    [1] link: "Home"
    [2] input:text: "Search"
    [12] editor: "What do you want to talk about?"
You act on elements BY THEIR NUMBER — this is reliable; matching by text is not.

  browser_navigate(url)          → Go to a URL. Automatically returns the numbered snapshot.
  browser_snapshot()             → Re-read the page. ALWAYS do this after a click that
                                   opens a dialog, expands a form, or changes the page.
  browser_click(text="5")        → Click element number 5.
  browser_type(element="12", content="...")        → Type into element number 12.
  browser_type(element="2", content="cats", submit=true) → Type, then press Enter.
  browser_key(key)               → Press a key (Enter, Tab, Escape).

THE LOOP — repeat until the task is done:
  1. browser_navigate to the right URL.
  2. Look at the numbered elements. The snapshot is split into FIELDS (type into) and
     BUTTONS/LINKS (click). MATCH each field by its label — e.g. to set a repository
     name, type into the field labelled "Repository name", NOT just element [1].
  3. Type into the right field(s) by number; click buttons by number.
  4. After any click that changes the page (opens a dialog/composer, submits a form),
     call browser_snapshot AGAIN — old numbers are invalid once the page changes, and a
     composer dialog adds NEW fields (the editor) that weren't there before.
  5. Finish by clicking the real submit button (e.g. "Create repository", "Post") by
     its number. Then snapshot once more to CONFIRM it worked before you reply.
You have up to 10 tool calls per command — use them to actually finish the job.

WRITING A POST / COMMENT (general pattern for any site):
  click the compose button → browser_snapshot (the editor is a NEW field now) →
  browser_type your full content into the editor → click the Post/Share button →
  browser_snapshot to confirm it appears. Do NOT stop right after opening the composer.

⚠ NEVER NARRATE INSTEAD OF ACTING. This is the most important rule.
  A spoken reply ENDS your turn. So if the task is not finished, DO NOT reply with
  words — call the next tool RIGHT NOW in the same turn. Phrases like "I'm taking a
  look", "let me gather", "I'm on the page and ready", "I'll now write it" are FORBIDDEN
  as a turn-ending reply while work remains. Either call a tool, or — only when truly
  done — give the short final reply. There is no in-between.

CONTINUING EARLIER WORK:
  You do NOT remember the previous page. If the user follows up about something you were
  doing ("I don't see the text", "now post it", "it didn't work", "try again"), your
  FIRST action MUST be browser_snapshot to read the LIVE page, then continue from there.

RULES:
  • Always write REAL, complete content — never placeholders like "content goes here".
    If web_search gives thin results, write from your own knowledge instead of stalling.
  • Pick the field whose LABEL matches what you're entering. Never type into a random
    field or a search box by mistake.
  • Use submit=true ONLY for a single search box. For multi-field forms / posts, type the
    content then CLICK the submit button — do not press Enter.
  • Element numbers change every time the page changes. Re-snapshot before acting again.
  • Never ask the user for a URL or what to type — figure it out. Only ask ONE short
    question if a required value (like a repo name) was genuinely not provided.
  • A task is NOT done after typing one field. Only say it's done after you clicked the
    final submit button AND a fresh snapshot confirms it. Never claim success otherwise.

━━ LOGIN PAGES ━━
If a page snapshot shows a sign-in or login form:
  STOP. Do NOT fill in any username or password.
  Say: "Please log into [site] in the Jarvis browser window, then ask me again."

━━ NATIVE WINDOWS APP SHORTCUTS ━━
Spotify    → run_shell_command("start spotify:")          ← always app, never spotify.com
Camera     → run_shell_command("start microsoft.windows.camera:")
Stopwatch  → run_shell_command("start ms-clock:Stopwatch")
Calculator → run_shell_command("start calc.exe")
Notepad    → run_shell_command("start notepad.exe")
Explorer   → run_shell_command("start explorer")
Task Mgr   → run_shell_command("taskmgr")
Shutdown   → run_shell_command("shutdown /s /t 1")
Restart    → run_shell_command("shutdown /r /t 1")

━━ URL PATTERNS ━━
YouTube search  : https://www.youtube.com/results?search_query=QUERY
Google search   : https://www.google.com/search?q=QUERY
Amazon India    : https://www.amazon.in/s?k=QUERY&sort=review-rank
Google Maps     : https://www.google.com/maps/dir/FROM/TO
Google News     : https://news.google.com/search?q=QUERY&hl=en-IN&gl=IN
Gmail inbox     : https://mail.google.com/mail/u/0/#inbox
Gmail compose   : https://mail.google.com/mail/?view=cm&to=EMAIL&su=SUBJECT&body=BODY
Google Meet     : https://meet.google.com/new
LinkedIn feed   : https://www.linkedin.com/feed/
Dominos India   : https://www.dominos.co.in/
Microsoft Store : https://apps.microsoft.com/search?query=QUERY

━━ USER'S PERSONAL INFO ━━
YouTube channel "Dietichen" : https://www.youtube.com/channel/UCNcVMyq5JyZ5V_TXN6KUgbA
YouTube channel "20xAI"     : https://www.youtube.com/channel/UCjz0PNtjFD_1haoyZbYaX-g
YouTube Studio "Dietichen"  : https://studio.youtube.com/channel/UCNcVMyq5JyZ5V_TXN6KUgbA
YouTube Studio "20xAI"      : https://studio.youtube.com/channel/UCjz0PNtjFD_1haoyZbYaX-g
Games available: breakout, car arcade, obstacle course, space invaders, football champs, path finder, tic tac toe"""


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  BRAIN  ─  Google Gemini (fully cloud, no local GPU)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ── Brain configuration ──────────────────────────────────────────────────────
# Brain: Google Gemini (free tier) via its OpenAI-compatible endpoint — our TOOLS
# are already in OpenAI format, so they work as-is. Multiple API keys are rotated
# automatically when one hits its daily rate limit (HTTP 429), and within each key
# we fall through across models. Nothing runs locally on the GPU.
#
def _load_dotenv() -> None:
    """Load KEY=VALUE pairs from a local, git-ignored `.env` file into the
    environment (without overriding anything already set). Lets you keep your
    Gemini keys in `.env`, e.g.  JARVIS_GEMINI_KEYS=key1,key2,key3"""
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env")
    try:
        with open(path, encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("#") or "=" not in line:
                    continue
                k, v = line.split("=", 1)
                os.environ.setdefault(k.strip(), v.strip().strip('"').strip("'"))
    except FileNotFoundError:
        pass


_load_dotenv()


# Keys are loaded from (in order of priority):
#   1. the JARVIS_GEMINI_KEYS env var / .env file (comma-separated), or
#   2. a local, GIT-IGNORED file `jarvis_keys.txt` next to this script — one key
#      per line, blank lines and #-comments ignored.
# Keys are NEVER hardcoded here, so this file is safe to commit/share publicly.
def _load_gemini_keys() -> list:
    env = [k.strip() for k in os.environ.get("JARVIS_GEMINI_KEYS", "").split(",") if k.strip()]
    if env:
        return env
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "jarvis_keys.txt")
    try:
        with open(path, encoding="utf-8") as f:
            return [ln.strip() for ln in f
                    if ln.strip() and not ln.strip().startswith("#")]
    except FileNotFoundError:
        return []


GEMINI_KEYS = _load_gemini_keys()
# Free tier is ~20 generate-content requests/DAY *per key, per MODEL*. An
# agentic task spends many requests (one per tool-loop round), so the daily
# quota — not the per-minute rate — is the real ceiling. Because the quota is
# per model, gemini-2.5-flash and gemini-2.5-flash-lite each have their OWN
# 20/day pool: we try flash first, then fall through to flash-lite (free extra
# budget). Override the whole list with JARVIS_GEMINI_MODELS (comma-separated)
# or a single model with JARVIS_GEMINI_MODEL.
GEMINI_MODELS = (
    [m.strip() for m in os.environ.get("JARVIS_GEMINI_MODELS", "").split(",") if m.strip()]
    or ([os.environ["JARVIS_GEMINI_MODEL"]] if os.environ.get("JARVIS_GEMINI_MODEL") else [])
    or ["gemini-2.5-flash", "gemini-2.5-flash-lite"]
)
GEMINI_MODEL = GEMINI_MODELS[0]   # primary (used for the startup reachability probe)
GEMINI_URL   = "https://generativelanguage.googleapis.com/v1beta/openai/chat/completions"

_gemini_key_idx = 0          # index of the key that last worked — rotation starts here
_gemini_last_model = GEMINI_MODEL   # model that actually served the last request


def _sanitize_assistant(msg: dict) -> dict:
    """Keep only the fields the chat APIs accept, so a message returned by one
    provider can be safely re-sent to either provider on the next round."""
    clean = {"role": "assistant", "content": msg.get("content")}
    tcs = msg.get("tool_calls")
    if tcs:
        clean["tool_calls"] = [
            {
                "id":   tc.get("id") or f"call_{i}",
                "type": "function",
                "function": {
                    "name":      tc["function"]["name"],
                    "arguments": tc["function"].get("arguments") or "{}",
                },
            }
            for i, tc in enumerate(tcs)
        ]
    return clean


def _gemini_request(messages, tools, temperature) -> dict:
    """Call Gemini, rotating through keys on rate-limit (429) and retrying
    transient busy errors (503). When every key is rate-limited on the primary
    model, fall through to the next model in GEMINI_MODELS — each model has its
    own separate daily free quota, so this is genuinely extra free budget.
    Returns the assistant message dict (OpenAI format). Raises RuntimeError if
    no key/model combination could serve it (the caller surfaces the error)."""
    global _gemini_key_idx, _gemini_last_model
    if not GEMINI_KEYS:
        raise RuntimeError("no Gemini keys configured")
    _usage_rollover()      # zero the meter if a new Pacific quota-day has started
    n = len(GEMINI_KEYS)
    last_err = "unknown"
    # Two retry rounds: if a whole sweep failed ONLY on transient per-minute
    # throttles (not the daily cap), wait briefly and try again before giving up.
    for _round in range(2):
        soft_only = True
        for model in GEMINI_MODELS:
            for offset in range(n):
                idx = (_gemini_key_idx + offset) % n
                key = GEMINI_KEYS[idx]
                body = {"model": model, "messages": messages, "temperature": temperature}
                if tools:
                    body["tools"] = tools
                for _attempt in range(3):
                    try:
                        r = requests.post(
                            GEMINI_URL,
                            headers={"Authorization": f"Bearer {key}",
                                     "Content-Type": "application/json"},
                            json=body, timeout=60,
                        )
                    except Exception as e:
                        # Transient network blip — don't kill the turn. Note it and let
                        # the loop retry / rotate; we only raise if EVERYTHING fails.
                        last_err = f"network error: {e}"
                        soft_only = False
                        time.sleep(1)
                        continue
                    if r.status_code == 200:
                        with _USAGE_LOCK:
                            _gemini_key_idx = idx          # remember the working key
                            _gemini_last_model = model     # …and the model that served it
                            _GEMINI_USAGE[(idx, model)] = _GEMINI_USAGE.get((idx, model), 0) + 1
                            _usage_save()
                        _hud_quota()
                        return r.json()["choices"][0]["message"]
                    if r.status_code == 503:
                        time.sleep(2)                  # model busy — retry same key
                        last_err = "503 busy"
                        continue
                    if r.status_code == 429:
                        # Distinguish the DAILY cap (key is done for the day) from a
                        # transient PER-MINUTE throttle (clears in ~seconds). Only the
                        # daily case should mark the key "full" on the quota meter.
                        if "perday" in r.text.lower().replace(" ", ""):
                            with _USAGE_LOCK:
                                _GEMINI_USAGE[(idx, model)] = 20   # genuinely tapped out today
                                _usage_save()
                            _hud_quota()
                            last_err = f"429 daily-limit (key {idx + 1}, {model})"
                        else:
                            soft_only = soft_only and True    # transient — don't poison meter
                            last_err = f"429 rate-throttle (key {idx + 1}, {model})"
                            break                              # rotate to a fresher key
                        soft_only = False
                        break                          # key tapped out — try the next key
                    last_err = f"HTTP {r.status_code}: {r.text[:120]}"
                    soft_only = False
                    break                              # 401/403/400 — try the next key
            # every key exhausted for this model → fall through to the next model
        if soft_only and _round == 0:
            time.sleep(15)   # whole sweep was only per-minute throttling — wait it out once
            continue
        break
    raise RuntimeError(f"all Gemini keys+models unavailable ({last_err})")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  VISION  ─  Gemini multimodal (Jarvis can see)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def _gemini_vision(prompt: str, image_b64: str, mime: str = "image/png") -> str:
    """Ask Gemini a question about an image. Uses the same key/model rotation as
    the text brain (vision works on the free tier). Returns the answer text."""
    messages = [{"role": "user", "content": [
        {"type": "text", "text": prompt},
        {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{image_b64}"}},
    ]}]
    msg = _gemini_request(messages, None, 0.2)   # tools omitted for pure vision
    return (msg.get("content") or "").strip() or "(no description returned)"


def _gemini_vision_multi(prompt: str, b64s: list, mime: str = "image/jpeg") -> str:
    """Ask Gemini about SEVERAL images in one call (e.g. a sequence of webcam
    frames over time). Returns the answer text."""
    content = [{"type": "text", "text": prompt}]
    for b in b64s:
        content.append({"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b}"}})
    msg = _gemini_request([{"role": "user", "content": content}], None, 0.2)
    return (msg.get("content") or "").strip() or "(no description returned)"


def _gemini_transcribe(audio_b64: str, fmt: str = "wav") -> str:
    """Transcribe audio with Gemini (free tier supports audio input). fmt is the
    container, e.g. 'wav' or 'mp3'. Returns the spoken words."""
    messages = [{"role": "user", "content": [
        {"type": "text", "text": "Transcribe this audio to text exactly. Output only the spoken words."},
        {"type": "input_audio", "input_audio": {"data": audio_b64, "format": fmt}},
    ]}]
    return (_gemini_request(messages, None, 0.1).get("content") or "").strip()


def _grab_screen_b64(max_w: int = 1280):
    """Capture the desktop, downscale, and return (base64_jpeg, mime)."""
    from PIL import ImageGrab
    img = ImageGrab.grab()
    if img.width > max_w:
        img = img.resize((max_w, int(img.height * max_w / img.width)))
    buf = io.BytesIO()
    img.convert("RGB").save(buf, "JPEG", quality=70)
    return base64.b64encode(buf.getvalue()).decode(), "image/jpeg"


def _vision_click(page, description: str) -> str:
    """Click an element by VISUAL description using Set-of-Marks prompting: every
    interactive element is tagged + drawn with its number on a screenshot, then
    Gemini picks WHICH number matches the description (reliable — it chooses among
    labelled candidates instead of guessing raw pixels). We click that element by
    its data-jarvis-id, so the click itself is exact. Great for icon buttons,
    images and controls whose text label is missing or unhelpful."""
    from PIL import Image, ImageDraw
    _index_interactive(page)   # tag every visible interactive element with data-jarvis-id
    rects = page.evaluate(
        "() => { const d = window.devicePixelRatio || 1;"
        " return [...document.querySelectorAll('[data-jarvis-id]')].map(e=>{"
        "  const r=e.getBoundingClientRect();"
        "  return {id:+e.getAttribute('data-jarvis-id'),"
        "          x:(r.x+r.width/2)*d, y:(r.y+r.height/2)*d,"
        "          on:(r.width>1&&r.height>1&&r.bottom>0&&r.top<innerHeight)};})"
        " .filter(o=>o.on); }"
    )
    if not rects:
        return "Nothing visible to click on this page."
    png = page.screenshot(type="png")
    img = Image.open(io.BytesIO(png)).convert("RGB")
    draw = ImageDraw.Draw(img)
    for o in rects:                       # draw numbered markers (Set-of-Marks)
        x, y = o["x"], o["y"]
        draw.ellipse([x - 14, y - 11, x + 14, y + 11], fill=(255, 40, 90))
        draw.text((x - 4 * len(str(o["id"])), y - 6), str(o["id"]), fill=(255, 255, 255))
    buf = io.BytesIO(); img.save(buf, "PNG")
    prompt = (
        "Each clickable element on this page is marked with a pink numbered badge. "
        f"Which badge number is on: \"{description}\"? "
        "Reply with ONLY JSON: {\"id\": <number>, \"found\": true} or {\"found\": false} if none match."
    )
    ans = _gemini_vision(prompt, base64.b64encode(buf.getvalue()).decode(), "image/png")
    mt = re.search(r"\{.*\}", ans, re.S)
    data = json.loads(mt.group(0)) if mt else {}
    if not data.get("found") or "id" not in data:
        return f"I looked but couldn't visually identify '{description}'. Try browser_snapshot for the numbered list."
    idx = int(data["id"])
    loc = page.locator(f"[data-jarvis-id='{idx}']")
    if loc.count() == 0:
        return f"Vision picked element [{idx}] but it's no longer on the page. Re-snapshot and retry."
    loc.first.click(timeout=5000)
    _settle(page)
    snap = execute_tool("browser_snapshot", {})
    return f"Saw and clicked '{description}' (element [{idx}]). Page now:\n\n{snap}"


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  LONG-TERM MEMORY  ─  free Gemini embeddings + local vector store
#  Jarvis embeds and stores durable facts; before each command the most
#  relevant ones are recalled by semantic similarity and fed back to the brain.
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
GEMINI_EMBED_URL   = "https://generativelanguage.googleapis.com/v1beta/openai/embeddings"
GEMINI_EMBED_MODEL = os.environ.get("JARVIS_EMBED_MODEL", "gemini-embedding-001")
GEMINI_EMBED_DIM   = 768   # compact, fast cosine; the model supports custom dims
_MEM_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "jarvis_memory.json")
_MEM = []   # [{"text":str, "ts":str, "vec":[float]}]


def _gemini_embed(texts):
    """Embed a list of strings via the free embeddings endpoint (key rotation).
    Returns a list of vectors, or [] on failure."""
    global _gemini_key_idx
    if not GEMINI_KEYS:
        return []
    n = len(GEMINI_KEYS)
    body = {"model": GEMINI_EMBED_MODEL, "input": texts, "dimensions": GEMINI_EMBED_DIM}
    for offset in range(n):
        idx = (_gemini_key_idx + offset) % n
        try:
            r = requests.post(
                GEMINI_EMBED_URL,
                headers={"Authorization": f"Bearer {GEMINI_KEYS[idx]}", "Content-Type": "application/json"},
                json=body, timeout=30,
            )
        except Exception:
            continue          # transient network blip on this key — try the next
        if r.status_code == 200:
            try:
                return [d["embedding"] for d in r.json()["data"]]
            except Exception:
                continue
        # 429 (quota) or any other error — rotate to the next key rather than give up
        continue
    return []


def _cosine(a, b):
    dot = s1 = s2 = 0.0
    for x, y in zip(a, b):
        dot += x * y; s1 += x * x; s2 += y * y
    return dot / (math.sqrt(s1) * math.sqrt(s2) + 1e-9)


def _mem_load():
    global _MEM
    try:
        with open(_MEM_PATH, encoding="utf-8") as f:
            _MEM = json.load(f)
    except Exception:
        _MEM = []


def _mem_save():
    try:
        with open(_MEM_PATH, "w", encoding="utf-8") as f:
            json.dump(_MEM, f)
    except Exception:
        pass


def _mem_add(fact: str) -> bool:
    """Embed and store a fact. Skips near-duplicates. Returns True on success."""
    vecs = _gemini_embed([fact])
    if not vecs:
        return False
    vec = vecs[0]
    for m in _MEM:                               # de-dupe very similar facts
        if m.get("vec") and _cosine(vec, m["vec"]) > 0.95:
            m["text"], m["vec"] = fact, vec
            m["ts"] = datetime.datetime.now().isoformat(timespec="seconds")
            _mem_save()
            _emit_memory_list()
            return True
    _MEM.append({"text": fact, "ts": datetime.datetime.now().isoformat(timespec="seconds"), "vec": vec})
    _mem_save()
    _emit_memory_list()
    return True


def _mem_search(query: str, k: int = 4, threshold: float = 0.30):
    """Return up to k (text, score) memories most similar to query, above threshold."""
    if not _MEM or not query.strip():
        return []
    qv = _gemini_embed([query])
    if not qv:
        return []
    qv = qv[0]
    scored = [(m["text"], _cosine(qv, m["vec"])) for m in _MEM if m.get("vec")]
    scored.sort(key=lambda x: x[1], reverse=True)
    return [(t, s) for t, s in scored[:k] if s >= threshold]


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  SYSTEM VITALS  ─  live CPU / RAM / GPU / battery telemetry → HUD
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def _gpu_stats():
    try:
        out = subprocess.run(
            ["nvidia-smi", "--query-gpu=utilization.gpu,memory.used,memory.total,temperature.gpu",
             "--format=csv,noheader,nounits"],
            capture_output=True, text=True, timeout=2,
            creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
        )
        u, mu, mt, t = [x.strip() for x in out.stdout.strip().split("\n")[0].split(",")]
        return {"util": float(u), "mem_used": float(mu), "mem_total": float(mt), "temp": float(t)}
    except Exception:
        return None


def _vitals() -> dict:
    import psutil
    v = {"cpu": round(psutil.cpu_percent()), "ram": round(psutil.virtual_memory().percent)}
    try:
        b = psutil.sensors_battery()
        if b is not None:
            v["batt"] = round(b.percent)
            v["plugged"] = bool(b.power_plugged)
    except Exception:
        pass
    g = _gpu_stats()
    if g:
        v["gpu"] = round(g["util"])
        v["gpu_mem"] = round(g["mem_used"] / g["mem_total"] * 100) if g["mem_total"] else 0
        v["gpu_temp"] = round(g["temp"])
    return v


def _vitals_loop() -> None:
    try:
        import psutil
        psutil.cpu_percent()      # prime the first reading
    except Exception:
        return
    while True:
        try:
            _hud_emit("vitals", **_vitals())
        except Exception:
            pass
        time.sleep(2)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  WEATHER  ─  free, no API key (Open-Meteo)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
_WMO = {0:"clear sky",1:"mainly clear",2:"partly cloudy",3:"overcast",45:"fog",48:"rime fog",
    51:"light drizzle",53:"drizzle",55:"dense drizzle",61:"light rain",63:"rain",65:"heavy rain",
    66:"freezing rain",71:"light snow",73:"snow",75:"heavy snow",77:"snow grains",80:"light showers",
    81:"showers",82:"violent showers",85:"snow showers",86:"heavy snow showers",95:"thunderstorm",
    96:"thunderstorm w/ hail",99:"severe thunderstorm"}


def _get_weather(location: str) -> str:
    try:
        g = requests.get("https://geocoding-api.open-meteo.com/v1/search",
                         params={"name": location, "count": 1}, timeout=10).json()
        if not g.get("results"):
            return f"Couldn't find a place called '{location}'."
        r = g["results"][0]
        w = requests.get("https://api.open-meteo.com/v1/forecast", params={
            "latitude": r["latitude"], "longitude": r["longitude"],
            "current": "temperature_2m,apparent_temperature,relative_humidity_2m,weather_code,wind_speed_10m",
            "timezone": "auto"}, timeout=10).json()
        c = w["current"]
        desc = _WMO.get(c.get("weather_code"), "unknown conditions")
        place = ", ".join(x for x in (r.get("name"), r.get("country")) if x)
        return (f"{place}: {c['temperature_2m']}°C and {desc}, feels like "
                f"{c['apparent_temperature']}°C, humidity {c['relative_humidity_2m']}%, "
                f"wind {c['wind_speed_10m']} km/h.")
    except Exception as e:
        return f"Weather error: {e}"


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  REMINDERS / TIMERS  ─  a background scheduler that speaks when due
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
_REM_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "jarvis_reminders.json")
_REMINDERS = []      # [{"text":str, "at":epoch_float}]
_REM_LOCK = threading.Lock()


def _rem_load():
    global _REMINDERS
    try:
        with open(_REM_PATH, encoding="utf-8") as f:
            _REMINDERS = json.load(f)
    except Exception:
        _REMINDERS = []


def _rem_save():
    try:
        with open(_REM_PATH, "w", encoding="utf-8") as f:
            json.dump(_REMINDERS, f)
    except Exception:
        pass


def _emit_reminders():
    with _REM_LOCK:
        rows = sorted(({"text": r["text"], "at": r["at"]} for r in _REMINDERS), key=lambda r: r["at"])
    _hud_emit("reminders", items=rows)


def _set_reminder(text: str, delay_seconds: int) -> str:
    when = time.time() + max(1, int(delay_seconds))
    with _REM_LOCK:
        _REMINDERS.append({"text": text, "at": when})
        _rem_save()
    _emit_reminders()
    mins = max(1, round(delay_seconds / 60))
    _hud_emit("toast", text=f"⏰ reminder set ({mins} min)", level="ok")
    return f"Reminder set. I'll remind you in about {mins} minute(s)."


def _reminder_loop() -> None:
    while True:
        now = time.time()
        due = []
        with _REM_LOCK:
            keep = []
            for r in _REMINDERS:
                (due if r["at"] <= now else keep).append(r)
            if due:
                _REMINDERS[:] = keep
                _rem_save()
        if due:
            _emit_reminders()
        for r in due:
            _hud_emit("toast", text=f"⏰ {r['text']}", level="warn")
            _away_note(f"a reminder fired ({r['text']})")
            try:
                _notify(f"Reminder: {r['text']}")   # speak + push to phone
            except Exception:
                pass
        time.sleep(3)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  CLIPBOARD + LOCAL FILE Q&A
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def _read_clipboard() -> str:
    try:
        import pyperclip
        return pyperclip.paste() or ""
    except Exception:
        try:
            out = subprocess.run(["powershell", "-NoProfile", "-Command", "Get-Clipboard"],
                                 capture_output=True, text=True, timeout=5)
            return out.stdout.strip()
        except Exception:
            return ""


def _read_file_text(path: str, limit: int = 14000) -> str:
    path = os.path.expanduser(path.strip().strip('"'))
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        from pypdf import PdfReader
        txt = "\n".join((pg.extract_text() or "") for pg in PdfReader(path).pages)
    elif ext == ".docx":
        import docx
        txt = "\n".join(p.text for p in docx.Document(path).paragraphs)
    else:                                   # txt / md / csv / code / etc.
        with open(path, encoding="utf-8", errors="ignore") as f:
            txt = f.read()
    return txt[:limit]


def _ask_file(path: str, question: str) -> str:
    try:
        text = _read_file_text(path)
    except FileNotFoundError:
        return f"I can't find the file: {path}"
    except Exception as e:
        return f"Couldn't read that file: {e}"
    if not text.strip():
        return "That file appears to be empty or unreadable (maybe a scanned PDF — try look_at_screen)."
    messages = [
        {"role": "system", "content": "Answer the user's question using ONLY the document below. Be concise and spoken-friendly. If the answer isn't in it, say so.\n\nDOCUMENT:\n" + text},
        {"role": "user", "content": question},
    ]
    try:
        return (_gemini_request(messages, None, 0.2).get("content") or "").strip()
    except Exception as e:
        return f"Couldn't analyse the file: {e}"


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  CALCULATOR  ─  safe arithmetic / math (local, instant, no quota)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
import ast as _ast
import operator as _op

_CALC_BIN = {_ast.Add: _op.add, _ast.Sub: _op.sub, _ast.Mult: _op.mul, _ast.Div: _op.truediv,
             _ast.Pow: _op.pow, _ast.Mod: _op.mod, _ast.FloorDiv: _op.floordiv}
_CALC_UN = {_ast.USub: _op.neg, _ast.UAdd: _op.pos}
_CALC_NAMES = {"pi": math.pi, "e": math.e, "tau": math.tau}
_CALC_FUNCS = {"sqrt": math.sqrt, "sin": math.sin, "cos": math.cos, "tan": math.tan,
               "log": math.log, "log10": math.log10, "exp": math.exp, "abs": abs,
               "round": round, "floor": math.floor, "ceil": math.ceil, "factorial": math.factorial,
               "radians": math.radians, "degrees": math.degrees, "min": min, "max": max}


def _calc_eval(node):
    if isinstance(node, _ast.Constant):
        if isinstance(node.value, (int, float)):
            return node.value
        raise ValueError("only numbers allowed")
    if isinstance(node, _ast.BinOp) and type(node.op) in _CALC_BIN:
        return _CALC_BIN[type(node.op)](_calc_eval(node.left), _calc_eval(node.right))
    if isinstance(node, _ast.UnaryOp) and type(node.op) in _CALC_UN:
        return _CALC_UN[type(node.op)](_calc_eval(node.operand))
    if isinstance(node, _ast.Name) and node.id in _CALC_NAMES:
        return _CALC_NAMES[node.id]
    if isinstance(node, _ast.Call) and isinstance(node.func, _ast.Name) and node.func.id in _CALC_FUNCS:
        return _CALC_FUNCS[node.func.id](*[_calc_eval(a) for a in node.args])
    raise ValueError("unsupported expression")


def _calculate(expr: str) -> str:
    try:
        val = _calc_eval(_ast.parse(expr.strip(), mode="eval").body)
        if isinstance(val, float) and val.is_integer():
            val = int(val)
        return f"{expr.strip()} = {val}"
    except Exception:
        return f"I couldn't compute '{expr}'. Try a plain arithmetic expression."


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  NEWS  ─  free headlines via Google News RSS (no key)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def _get_news(topic: str = "", n: int = 6) -> list:
    import xml.etree.ElementTree as ET
    if topic.strip():
        url = "https://news.google.com/rss/search?q=" + urllib.parse.quote(topic) + "&hl=en-US&gl=US&ceid=US:en"
    else:
        url = "https://news.google.com/rss?hl=en-US&gl=US&ceid=US:en"
    try:
        xml = requests.get(url, timeout=10, headers={"User-Agent": "Mozilla/5.0"}).text
        root = ET.fromstring(xml)
        items = root.findall(".//item")
        return [it.findtext("title", "").strip() for it in items[:n] if it.findtext("title")]
    except Exception:
        return []


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  DAILY BRIEFING  ─  greeting + weather + headlines + pending reminders
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def _compose_briefing(location: str = "", topic: str = "") -> str:
    hour = datetime.datetime.now().hour
    greet = "Good morning" if hour < 12 else "Good afternoon" if hour < 18 else "Good evening"
    parts = [f"{greet}. It's {datetime.datetime.now().strftime('%A, %B %d, %I:%M %p')}."]
    if location:
        w = _get_weather(location)
        if w and not w.lower().startswith(("weather error", "couldn't")):
            parts.append("Weather: " + w)
    heads = _get_news(topic, n=4)
    if heads:
        parts.append("Top headlines: " + "; ".join(heads) + ".")
    with _REM_LOCK:
        pending = len(_REMINDERS)
    if pending:
        parts.append(f"You have {pending} reminder(s) pending.")
    return " ".join(parts)


# Proactive routines: optional jarvis_routines.json like
#   [{"time": "08:00", "location": "Vellore", "topic": "AI"}]
# fires a spoken briefing once when the local clock reaches that HH:MM.
_ROUTINES_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "jarvis_routines.json")


def _routine_loop() -> None:
    fired_today = set()
    last_day = None
    while True:
        try:
            now = datetime.datetime.now()
            if now.day != last_day:
                fired_today.clear(); last_day = now.day      # reset each day
            try:
                with open(_ROUTINES_PATH, encoding="utf-8") as f:
                    routines = json.load(f)
            except Exception:
                routines = []
            hhmm = now.strftime("%H:%M")
            for r in routines:
                key = (r.get("time"), r.get("location", ""), r.get("topic", ""))  # stable id, not list index
                if r.get("time") == hhmm and key not in fired_today:
                    fired_today.add(key)
                    _hud_emit("toast", text="🌅 daily briefing", level="info")
                    speak(_compose_briefing(r.get("location", ""), r.get("topic", "")))
        except Exception:
            pass
        time.sleep(20)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  PLAN  ─  the model declares a step plan; the HUD shows live progress
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
_PLAN = {"steps": [], "done": 0}


def _set_plan(steps) -> str:
    steps = [str(s).strip() for s in (steps or []) if str(s).strip()][:8]
    _PLAN["steps"], _PLAN["done"] = steps, 0
    _hud_emit("plan", steps=steps, done=0)
    return "Plan set." if steps else "Empty plan."


def _complete_step(index) -> str:
    try:
        i = int(index)
    except Exception:
        i = _PLAN["done"]
    _PLAN["done"] = max(_PLAN["done"], i + 1)
    _hud_emit("plan", steps=_PLAN["steps"], done=_PLAN["done"])
    return "Step marked done."


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  WRITE FILES  ·  CURRENCY  ·  MEDIA KEYS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def _write_file(path: str, content: str) -> str:
    """Save text/markdown/code (or a .docx) to disk. Creates parent dirs."""
    path = os.path.expanduser(path.strip().strip('"'))
    try:
        d = os.path.dirname(path)
        if d:
            os.makedirs(d, exist_ok=True)
        if path.lower().endswith(".docx"):
            import docx
            doc = docx.Document()
            for line in content.split("\n"):
                doc.add_paragraph(line)
            doc.save(path)
        else:
            with open(path, "w", encoding="utf-8") as f:
                f.write(content)
        return f"Saved {len(content)} characters to {path}."
    except Exception as e:
        return f"Couldn't write the file: {e}"


def _convert_currency(amount: float, frm: str, to: str) -> str:
    frm, to = frm.upper().strip(), to.upper().strip()
    try:
        data = requests.get(f"https://open.er-api.com/v6/latest/{frm}", timeout=10).json()
        rate = (data.get("rates") or {}).get(to)
        if not rate:
            return f"Couldn't get a {frm}->{to} rate."
        return f"{amount} {frm} = {round(amount * rate, 2)} {to} (rate {round(rate, 4)})."
    except Exception as e:
        return f"Currency error: {e}"


def _fetch_page_text(url: str, limit: int = 12000) -> str:
    """Fetch a URL and return readable plain text (no LLM call). '' on failure."""
    if not url.startswith("http"):
        url = "https://" + url
    try:
        html = requests.get(url, timeout=15, headers={"User-Agent": "Mozilla/5.0"}).text
    except Exception:
        return ""
    html = re.sub(r"(?is)<(script|style|noscript|svg|header|footer|nav)[^>]*>.*?</\1>", " ", html)
    text = re.sub(r"(?s)<[^>]+>", " ", html)
    text = re.sub(r"&[a-z]+;", " ", text)
    return re.sub(r"\s+", " ", text).strip()[:limit]


def _read_url(url: str, question: str = "") -> str:
    """Fetch a web page, extract readable text, and summarise/answer with Gemini."""
    text = _fetch_page_text(url)
    if not text:
        return "Couldn't fetch or read that page."
    task = question.strip() or "Summarise this page in 3-4 spoken-friendly sentences."
    messages = [
        {"role": "system", "content": "Use ONLY the page text below. Be concise and spoken-friendly.\n\nPAGE:\n" + text},
        {"role": "user", "content": task},
    ]
    try:
        return (_gemini_request(messages, None, 0.3).get("content") or "").strip()
    except Exception as e:
        return f"Couldn't analyse the page: {e}"


_COIN_IDS = {"btc": "bitcoin", "bitcoin": "bitcoin", "eth": "ethereum", "ethereum": "ethereum",
             "sol": "solana", "solana": "solana", "doge": "dogecoin", "dogecoin": "dogecoin",
             "ada": "cardano", "xrp": "ripple", "bnb": "binancecoin", "matic": "matic-network",
             "ltc": "litecoin", "dot": "polkadot", "shib": "shiba-inu"}


def _get_crypto_price(coin: str, currency: str = "usd") -> str:
    cid = _COIN_IDS.get(coin.lower().strip(), coin.lower().strip().replace(" ", "-"))
    cur = (currency or "usd").lower().strip()
    try:
        d = requests.get("https://api.coingecko.com/api/v3/simple/price",
                         params={"ids": cid, "vs_currencies": cur, "include_24hr_change": "true"},
                         timeout=10).json()
        if cid not in d:
            return f"Couldn't find a crypto called '{coin}'."
        price = d[cid][cur]
        chg = d[cid].get(f"{cur}_24h_change")
        extra = f", {chg:+.1f}% in 24h" if isinstance(chg, (int, float)) else ""
        return f"{cid.capitalize()} is {price:,} {cur.upper()}{extra}."
    except Exception as e:
        return f"Crypto price error: {e}"


def _define_word(word: str) -> str:
    word = word.strip()
    try:
        d = requests.get(f"https://api.dictionaryapi.dev/api/v2/entries/en/{urllib.parse.quote(word)}",
                         timeout=10).json()
        if not isinstance(d, list):
            return f"I couldn't find a definition for '{word}'."
        out = []
        for meaning in d[0].get("meanings", [])[:3]:
            pos = meaning.get("partOfSpeech", "")
            defs = meaning.get("definitions", [])
            if defs:
                out.append(f"({pos}) {defs[0].get('definition','')}")
        return f"{word}: " + " ".join(out) if out else f"No definition found for '{word}'."
    except Exception as e:
        return f"Dictionary error: {e}"


_UNIT_BASE = {  # convert everything to a base unit per dimension
    # length → metres
    "mm": 0.001, "cm": 0.01, "m": 1, "km": 1000, "inch": 0.0254, "in": 0.0254,
    "ft": 0.3048, "foot": 0.3048, "feet": 0.3048, "yard": 0.9144, "yd": 0.9144,
    "mile": 1609.344, "mi": 1609.344,
    # mass → grams
    "mg": 0.001, "g": 1, "gram": 1, "kg": 1000, "lb": 453.592, "lbs": 453.592,
    "pound": 453.592, "oz": 28.3495, "ounce": 28.3495, "ton": 1_000_000,
    # volume → litres
    "ml": 0.001, "l": 1, "litre": 1, "liter": 1, "gal": 3.78541, "gallon": 3.78541,
    "cup": 0.236588, "pint": 0.473176,
    # speed → m/s
    "mps": 1, "kmh": 0.277778, "kph": 0.277778, "mph": 0.44704,
    # data → bytes
    "b": 1, "kb": 1024, "mb": 1024**2, "gb": 1024**3, "tb": 1024**4,
}
_UNIT_DIM = {}
for _grp in [["mm","cm","m","km","inch","in","ft","foot","feet","yard","yd","mile","mi"],
             ["mg","g","gram","kg","lb","lbs","pound","oz","ounce","ton"],
             ["ml","l","litre","liter","gal","gallon","cup","pint"],
             ["mps","kmh","kph","mph"], ["b","kb","mb","gb","tb"]]:
    for _u in _grp:
        _UNIT_DIM[_u] = _grp[0]


_UNIT_ALIAS = {"miles": "mi", "mile": "mi", "kilometers": "km", "kilometres": "km",
               "kilometer": "km", "kilometre": "km", "meters": "m", "metres": "m",
               "metre": "m", "meter": "m", "centimeters": "cm", "centimetres": "cm",
               "millimeters": "mm", "inches": "in", "feet": "ft", "yards": "yd",
               "pounds": "lb", "kilograms": "kg", "kilogram": "kg", "grams": "g",
               "ounces": "oz", "litres": "l", "liters": "l", "gallons": "gal",
               "cups": "cup", "pints": "pint", "celsius": "c", "fahrenheit": "f",
               "kelvin": "k", "centigrade": "c"}


def _unit_norm(u: str) -> str:
    u = u.lower().strip()
    if u in _UNIT_ALIAS:
        return _UNIT_ALIAS[u]
    if u not in _UNIT_BASE and u.endswith("s") and u[:-1] in _UNIT_BASE:
        return u[:-1]
    return u


def _convert_units(value: float, frm: str, to: str) -> str:
    f, t = _unit_norm(frm), _unit_norm(to)
    temp = {"c", "celsius", "f", "fahrenheit", "k", "kelvin"}
    if f in temp and t in temp:
        c = value if f in ("c", "celsius") else (value - 32) * 5 / 9 if f in ("f", "fahrenheit") else value - 273.15
        out = c if t in ("c", "celsius") else c * 9 / 5 + 32 if t in ("f", "fahrenheit") else c + 273.15
        return f"{value}°{f[0].upper()} = {round(out, 2)}°{t[0].upper()}."
    if f in _UNIT_BASE and t in _UNIT_BASE and _UNIT_DIM.get(f) == _UNIT_DIM.get(t):
        out = value * _UNIT_BASE[f] / _UNIT_BASE[t]
        return f"{value} {frm} = {round(out, 4)} {to}."
    return f"I can't convert {frm} to {to} (incompatible or unknown units)."


def _media_control(action: str) -> str:
    keymap = {"play": "play/pause media", "pause": "play/pause media",
              "playpause": "play/pause media", "next": "next track",
              "previous": "previous track", "prev": "previous track",
              "stop": "stop media", "volup": "volume up", "voldown": "volume down",
              "mute": "volume mute"}
    key = keymap.get(action.lower().replace(" ", ""))
    if not key:
        return f"Unknown media action '{action}'."
    try:
        import keyboard
        keyboard.send(key)
        return f"Sent media key: {action}."
    except Exception as e:
        return f"Couldn't send media key: {e}"


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  AUTONOMOUS RESEARCH  ─  multi-source web research → cited synthesis
#  Searches + page fetches are FREE (DuckDuckGo + plain HTTP); only the final
#  synthesis is a single Gemini call, so a deep dive costs ~1 request.
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def _deep_research(topic: str, depth: int = 3) -> str:
    depth = max(1, min(5, int(depth)))
    _set_plan([f"Search: {topic[:30]}", "Read sources", "Synthesise"])
    _hud_emit("toast", text="🔬 researching…", level="info")
    # 1) gather candidate sources (free)
    urls = []
    try:
        from ddgs import DDGS
        with DDGS() as ddgs:
            for r in ddgs.text(topic, max_results=depth + 3):
                if r.get("href"):
                    urls.append((r["title"], r["href"]))
    except Exception as e:
        return f"Research search failed: {e}"
    _complete_step(0)
    # 2) read the top pages (free)
    sources, used = [], []
    for title, url in urls:
        if len(sources) >= depth:
            break
        txt = _fetch_page_text(url, limit=4000)
        if len(txt) > 300:
            sources.append(f"SOURCE [{len(sources)+1}] {title} — {url}\n{txt}")
            used.append(f"[{len(sources)}] {url}")
            _hud_emit("tool", name="read_url", summary=title[:48], status="ok")
    _complete_step(1)
    if not sources:
        return "I couldn't read enough sources to research that."
    # 3) one synthesis call
    prompt = ("Research the user's topic using the sources below. Give a clear, well-structured "
              "answer (5-8 sentences) that synthesises across sources and notes any disagreement. "
              "End with a one-line 'Sources:' list of the numbers used.\n\nTOPIC: " + topic +
              "\n\n" + "\n\n".join(sources)[:28000])
    try:
        ans = (_gemini_request([{"role": "user", "content": prompt}], None, 0.3).get("content") or "").strip()
    except Exception as e:
        return f"Research synthesis failed: {e}"
    _complete_step(2)
    return ans + "\n\nSources:\n" + "\n".join(used)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  PROACTIVE WATCH MODE  ─  periodic vision glances that offer help
#  OFF by default (each glance spends 1 Gemini request). User toggles it on.
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
_WATCH = {"on": False, "interval": 90}


def _watch_glance() -> str:
    """One proactive look at the screen. Returns a suggestion, or '' if nothing useful."""
    try:
        b64, mime = _grab_screen_b64()
    except Exception:
        return ""
    prompt = ("You are a proactive assistant glancing at the user's screen. If they appear "
              "stuck, have an error, or there's something genuinely helpful you could offer, "
              "reply with ONE short spoken suggestion (max 20 words). If nothing is worth "
              "interrupting them for, reply with exactly: NONE")
    try:
        ans = _gemini_vision(prompt, b64, mime).strip()
    except Exception:
        return ""
    return "" if ans.upper().startswith("NONE") or len(ans) < 4 else ans


def _watch_loop() -> None:
    while True:
        if _WATCH["on"]:
            tip = _watch_glance()
            if tip:
                _hud_emit("toast", text="💡 " + tip[:60], level="info")
                try:
                    speak(tip)
                except Exception:
                    pass
        time.sleep(max(30, int(_WATCH["interval"])) if _WATCH["on"] else 5)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  MEMORY DRAWER + FORGET
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def _emit_memory_list():
    _hud_emit("memory_list", items=[m["text"] for m in _MEM][-60:], count=len(_MEM))


def _forget(query: str) -> str:
    hits = _mem_search(query, k=1, threshold=0.35)
    if not hits:
        return "I couldn't find a matching memory to forget."
    target = hits[0][0]
    before = len(_MEM)
    _MEM[:] = [m for m in _MEM if m["text"] != target]
    if len(_MEM) < before:
        _mem_save()
        _emit_memory_list()
        _hud_emit("memory", count=len(_MEM))
        return f"Forgotten: {target}"
    return "I couldn't remove that memory."


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  UNIVERSAL DESKTOP CONTROL  ─  operate ANY Windows app by sight
#  Set-of-Marks on real UI-Automation rectangles (reliable, not pixel-guessing):
#  number every visible control of the foreground window on a screenshot, let
#  Gemini pick the one matching the description, then act on its real centre
#  with pyautogui. Works in any application, not just the browser.
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def _desktop_marks_data(window_title: str = "") -> dict:
    """Enumerate the foreground window's visible controls (id, centre, label) using
    UI Automation. Runs in a dedicated COM-initialised thread because pywinauto's
    comtypes clashes with the COM apartment pygame/Playwright set on the main
    thread ('Cannot change thread mode'). Returns {title, marks} or {error}."""
    out = {}

    def _work():
        try:
            try:
                import comtypes
                comtypes.CoInitialize()
            except Exception:
                pass
            import ctypes
            from pywinauto import Desktop
            if window_title:
                try:
                    Desktop(backend="uia").window(title_re=f".*{re.escape(window_title)}.*").set_focus()
                    time.sleep(0.5)
                except Exception:
                    pass
            hwnd = ctypes.windll.user32.GetForegroundWindow()
            win = Desktop(backend="uia").window(handle=hwnd)
            out["title"] = win.window_text()
            marks = []
            for c in win.descendants():
                try:
                    if not c.is_visible():
                        continue
                    r = c.rectangle()
                    w, h = r.right - r.left, r.bottom - r.top
                    if w < 6 or h < 6 or w > 1900 or h > 1050:
                        continue
                    label = (c.window_text() or c.element_info.control_type or "")[:40]
                    marks.append((r.left + w // 2, r.top + h // 2, label))
                except Exception:
                    continue
                if len(marks) >= 70:
                    break
            out["marks"] = marks
        except Exception as e:
            out["error"] = str(e)

    t = threading.Thread(target=_work, daemon=True)
    t.start()
    t.join(timeout=20)
    return out


def _desktop_pick(description: str, window_title: str = ""):
    """Return {found, id, x, y, n, title} for the control matching *description*
    on the foreground window, using Set-of-Marks vision."""
    from PIL import ImageGrab, ImageDraw
    data = _desktop_marks_data(window_title)
    if data.get("error"):
        return {"found": False, "error": data["error"]}
    title = data.get("title", "")
    marks = data.get("marks", [])
    if not marks:
        return {"found": False, "title": title, "n": 0}
    img = ImageGrab.grab().convert("RGB")
    draw = ImageDraw.Draw(img)
    for i, (x, y, _lab) in enumerate(marks):
        draw.ellipse([x - 13, y - 10, x + 13, y + 10], fill=(255, 40, 90))
        draw.text((x - 4 * len(str(i)), y - 6), str(i), fill=(255, 255, 255))
    buf = io.BytesIO(); img.convert("RGB").save(buf, "JPEG", quality=70)
    prompt = (f"Each UI control is marked with a pink numbered badge. Which badge number is on: "
              f"\"{description}\"? Reply ONLY JSON: {{\"id\": <n>, \"found\": true}} or {{\"found\": false}}.")
    ans = _gemini_vision(prompt, base64.b64encode(buf.getvalue()).decode(), "image/jpeg")
    mt = re.search(r"\{.*\}", ans, re.S)
    data = json.loads(mt.group(0)) if mt else {}
    if not data.get("found") or "id" not in data or not (0 <= int(data["id"]) < len(marks)):
        return {"found": False, "title": title, "n": len(marks)}
    x, y, lab = marks[int(data["id"])]
    return {"found": True, "id": int(data["id"]), "x": x, "y": y, "n": len(marks), "label": lab, "title": title}


def _computer_control(description: str, action: str = "click", text: str = "", window: str = "") -> str:
    try:
        import pyautogui
    except Exception:
        return "Desktop control needs pyautogui (pip install pyautogui)."
    pick = _desktop_pick(description, window)
    if not pick.get("found"):
        return (f"I couldn't visually find '{description}' in '{pick.get('title','the active window')}' "
                f"(scanned {pick.get('n',0)} controls).")
    x, y = pick["x"], pick["y"]
    act = (action or "click").lower()
    try:
        pyautogui.moveTo(x, y, duration=0.15)
        if act in ("click", "left", ""):
            pyautogui.click(x, y)
        elif act in ("double", "double_click", "doubleclick"):
            pyautogui.doubleClick(x, y)
        elif act in ("right", "right_click", "rightclick"):
            pyautogui.rightClick(x, y)
        elif act == "type":
            pyautogui.click(x, y); time.sleep(0.2)
            pyautogui.typewrite(text, interval=0.02)
        else:
            pyautogui.click(x, y)
        _hud_emit("tool", name="computer_control", summary=f"{act}: {pick.get('label','')[:30]}", status="ok")
        return f"Did '{act}' on '{description}' (control [{pick['id']}] '{pick.get('label','')}') in '{pick['title']}'."
    except Exception as e:
        return f"Desktop action error: {e}"


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  SELF-AUTHORED SKILLS  ─  Jarvis writes & hot-loads new tools at runtime
#  Saved to jarvis_skills/<name>.py (reviewable). Each defines run(args)->str.
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
_SKILLS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "jarvis_skills")
_SKILLS = {}   # name -> {"func": callable, "description": str}


def _load_one_skill(path: str) -> bool:
    import importlib.util as _il
    name = os.path.splitext(os.path.basename(path))[0]
    try:
        spec = _il.spec_from_file_location(f"jarvis_skill_{name}", path)
        mod = _il.module_from_spec(spec)
        spec.loader.exec_module(mod)
        if not hasattr(mod, "run"):
            return False
        desc = getattr(mod, "DESCRIPTION", f"Custom skill '{name}'.")
        _SKILLS[name] = {"func": mod.run, "description": desc}
        _register_skill_tool(name, desc)
        return True
    except Exception as e:
        print(f"   ⚠ skill '{name}' failed to load: {e}")
        return False


def _register_skill_tool(name: str, description: str):
    if any(t["function"]["name"] == name for t in TOOLS):
        return
    TOOLS.append({"type": "function", "function": {
        "name": name,
        "description": description + " (a skill Jarvis taught itself)",
        "parameters": {"type": "object", "properties": {
            "input": {"type": "string", "description": "Input for the skill."}}},
    }})


def _load_skills():
    if not os.path.isdir(_SKILLS_DIR):
        return
    for fn in os.listdir(_SKILLS_DIR):
        if fn.endswith(".py"):
            _load_one_skill(os.path.join(_SKILLS_DIR, fn))


def _teach_skill(name: str, description: str, code: str) -> str:
    name = re.sub(r"[^a-z0-9_]", "", name.lower().strip())
    if not name:
        return "Invalid skill name."
    if "def run(" not in code:
        return "Skill code must define a function: def run(args): ... returning a string."
    os.makedirs(_SKILLS_DIR, exist_ok=True)
    path = os.path.join(_SKILLS_DIR, f"{name}.py")
    content = f"DESCRIPTION = {description!r}\n\n{code}\n"
    try:
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
    except Exception as e:
        return f"Couldn't save skill: {e}"
    if _load_one_skill(path):
        _hud_emit("toast", text=f"🧩 learned skill: {name}", level="ok")
        return f"Learned a new skill '{name}'. I can use it from now on."
    return f"Saved '{name}' but it failed to load — the code may have an error."


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  AUTONOMOUS MISSION MODE  ─  plan → act → self-verify → loop until done
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
_MISSION_SYS = (
    "You are Jarvis in AUTONOMOUS MISSION mode. You are given a GOAL and must achieve it "
    "end-to-end without further user input. Rules:\n"
    "1. FIRST call set_plan with the concrete steps.\n"
    "2. Execute steps using your tools; call complete_step after each.\n"
    "3. VERIFY your own work — after web/desktop actions use look_at_page or look_at_screen "
    "to confirm the real outcome before believing a step succeeded. If a step failed, adapt "
    "and retry a different way.\n"
    "4. Only when the goal is genuinely verified complete, reply with a short final summary "
    "starting with 'MISSION COMPLETE:'. Do not claim success you haven't verified."
)


def _mission(goal: str, max_rounds: int = 12) -> str:
    messages = [{"role": "system", "content": _MISSION_SYS},
                {"role": "user", "content": f"GOAL: {goal}"}]
    _hud_emit("toast", text="🚀 mission started", level="info")
    mission_tools = [t for t in TOOLS if t["function"]["name"] != "mission"]
    final = "Mission ended."
    for _ in range(max_rounds):
        if _ABORT.is_set():
            return "Mission stopped."
        try:
            assistant, _label = brain_chat(messages, mission_tools, 0.2)
        except Exception as e:
            return f"Mission error: {e}"
        messages.append(assistant)
        content = (assistant.get("content") or "").strip()
        if content:
            final = content
        calls = assistant.get("tool_calls")
        if not calls:
            break                                  # model is done talking
        for tc in calls:
            nm = tc["function"]["name"]
            try:
                a = json.loads(tc["function"].get("arguments") or "{}")
            except Exception:
                a = {}
            res = execute_tool(nm, a)
            messages.append({"role": "tool", "tool_call_id": tc.get("id", "call_0"), "content": str(res)[:2500]})
        if "MISSION COMPLETE" in content.upper():
            break
    _hud_emit("toast", text="🏁 mission finished", level="ok")
    return final


def brain_chat(messages, tools, temperature):
    """Return (assistant_message_dict, brain_label). Fully Gemini — rotates across
    keys and models. If every key/model is unavailable it raises, and the caller
    reports the error (no local model runs, so nothing touches the GPU)."""
    msg = _gemini_request(messages, tools, temperature)
    _hud_emit("brain", provider="gemini", model=_gemini_last_model, key=_gemini_key_idx + 1)
    return _sanitize_assistant(msg), f"Gemini ({_gemini_last_model}, key {_gemini_key_idx + 1})"


def process_command(user_input: str, history: list) -> tuple[str, list]:
    history.append({"role": "user", "content": user_input})

    system = SYSTEM_PROMPT
    # ── Auto-recall: surface relevant long-term memories for this command ──────
    try:
        hits = _mem_search(user_input, k=4)
        if hits:
            recalled = "\n".join(f"- {t}" for t, _ in hits)
            system += ("\n\n━━ THINGS YOU REMEMBER ABOUT THIS USER (use if relevant) ━━\n"
                       + recalled)
            print(f"   🧠 recalled {len(hits)} memory(ies)")
            _hud_emit("toast", text=f"🧠 recalled {len(hits)} memory(ies)", level="info")
            _hud_emit("memory", count=len(_MEM), recalled=len(hits))
    except Exception:
        pass

    messages = [{"role": "system", "content": system}] + history[-20:]

    final = "Done."
    label_shown = False
    _hud_emit("state", state="thinking", sub="reasoning")
    for _ in range(10):  # max 10 tool rounds per command
        if _ABORT.is_set():
            history.append({"role": "assistant", "content": "Okay, stopped."})
            return "__ABORTED__", history

        try:
            _hud_emit("state", state="thinking", sub="reasoning")
            assistant, label = brain_chat(messages, TOOLS, 0.15)
        except Exception as e:
            err = f"AI error: {e}"
            history.append({"role": "assistant", "content": err})
            return err, history

        if _ABORT.is_set():           # user hit ESC while the model was thinking
            history.append({"role": "assistant", "content": "Okay, stopped."})
            return "__ABORTED__", history

        if not label_shown:
            print(f"   🧠 {label}")
            label_shown = True

        messages.append(assistant)

        content = (assistant.get("content") or "").strip()
        if content:
            final = content

        tool_calls = assistant.get("tool_calls")
        if not tool_calls:
            break

        for tc in tool_calls:
            if _ABORT.is_set():       # stop before running any further tools
                history.append({"role": "assistant", "content": "Okay, stopped."})
                return "__ABORTED__", history
            name = tc["function"]["name"]
            try:
                args = json.loads(tc["function"].get("arguments") or "{}")
            except Exception:
                args = {}
            _hud_emit("state", state="acting", sub=name.replace("_", " "))
            _hud_emit("tool", name=name, summary=_tool_summary(name, args), status="run")
            result = execute_tool(name, args)
            _status, _detail = _tool_outcome(result)
            _hud_emit("tool", name=name, status=_status, detail=_detail)
            preview = str(result)
            if len(preview) > 900:
                preview = preview[:900] + " …(truncated)"
            print(f"   ↳ {preview}")
            messages.append({
                "role": "tool",
                "tool_call_id": tc.get("id", "call_0"),
                "content": str(result),
            })

    history.append({"role": "assistant", "content": final})
    return final, history


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  VISUAL PRESENCE  ─  Jarvis can SEE you (webcam)
#  Always-on awareness is LOCAL & FREE (OpenCV face/profile detection): it knows
#  when you arrive/leave, watches your screen-time, and streams a live thumbnail
#  to the HUD. Rich understanding (what you're holding, posture, mood) uses
#  Gemini vision on demand, so the camera never burns quota by just watching.
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
JARVIS_CAM_INDEX = int(os.environ.get("JARVIS_CAM_INDEX", "0") or 0)
_cam = None
_cam_lock = threading.Lock()
_cascades = None
_PRESENCE = {"on": os.environ.get("JARVIS_PRESENCE", "1") != "0",
             "present": False, "since": 0.0, "last_seen": 0.0,
             "away_since": 0.0, "last_break": 0.0,
             "desk_seconds": 0.0, "away_count": 0, "tick": 0.0}
# Optional Gemini posture check every N minutes while present (0 = off; costs quota)
_POSTURE_EVERY = int(os.environ.get("JARVIS_POSTURE_MINUTES", "0") or 0)
# Away-actions (opt-in): pause media when you leave / resume on return; lock the PC
_AWAY_PAUSE = os.environ.get("JARVIS_AWAY_PAUSE", "0") != "0"
_AWAY_LOCK_SEC = int(os.environ.get("JARVIS_AWAY_LOCK_SECONDS", "0") or 0)   # 0 = never
_AWAY_ALERT = os.environ.get("JARVIS_AWAY_ALERT", "0") != "0"   # ping phone on desk movement while away
_AWAY_LOG = []        # notable events while you're away → recapped on return
_CAM_RING = []        # in-memory ring of recent (ts, jpeg_b64) for camera recall (local only)
_paused_for_away = False
_locked_for_away = False


def _away_note(text: str) -> None:
    """Record something that happened while the user is away (for the recap)."""
    if not _PRESENCE.get("present", True):
        _AWAY_LOG.append((time.time(), text))


def _cascades_load():
    global _cascades
    if _cascades is None:
        try:
            import cv2
            H = cv2.data.haarcascades
            _cascades = [cv2.CascadeClassifier(H + "haarcascade_frontalface_default.xml"),
                         cv2.CascadeClassifier(H + "haarcascade_profileface.xml")]
        except Exception:
            _cascades = []
    return _cascades


def _cam_open():
    global _cam
    import cv2
    if _cam is None or not _cam.isOpened():
        _cam = cv2.VideoCapture(JARVIS_CAM_INDEX)
        for _ in range(5):           # warm-up frames (first ones are often blank)
            _cam.read(); time.sleep(0.05)
    return _cam


def _cam_release():
    global _cam
    with _cam_lock:          # don't release while another thread is mid-read
        try:
            if _cam is not None:
                _cam.release()
        except Exception:
            pass
        _cam = None


def _cam_frame():
    """Grab a fresh BGR frame (thread-safe). Returns ndarray or None."""
    with _cam_lock:
        try:
            cap = _cam_open()
            f = None
            for _ in range(3):       # flush buffered/stale frames → newest
                ok, fr = cap.read()
                if ok:
                    f = fr
            return f
        except Exception:
            return None


def _face_present(frame) -> bool:
    """LOCAL, free presence check — frontal + profile (both directions) for the
    45° angle. No network, no GPU model."""
    try:
        import cv2
        gray = cv2.equalizeHist(cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY))
    except Exception:
        return False
    cs = _cascades_load()
    if not cs:
        return False
    if len(cs[0].detectMultiScale(gray, 1.1, 4)) > 0:
        return True
    if len(cs) > 1:
        if len(cs[1].detectMultiScale(gray, 1.1, 4)) > 0:
            return True
        try:
            if len(cs[1].detectMultiScale(cv2.flip(gray, 1), 1.1, 4)) > 0:   # other profile
                return True
        except Exception:
            pass
    return False


def _cam_jpeg_b64(frame, max_w: int = 640, q: int = 75) -> str:
    import cv2
    h, w = frame.shape[:2]
    if w > max_w:
        frame = cv2.resize(frame, (max_w, int(h * max_w / w)))
    ok, buf = cv2.imencode(".jpg", frame, [cv2.IMWRITE_JPEG_QUALITY, q])
    return base64.b64encode(buf).decode()


def _cam_vision(question: str) -> str:
    frame = _cam_frame()
    if frame is None:
        return "I couldn't get a frame from the webcam (is it connected / set JARVIS_CAM_INDEX?)."
    return _gemini_vision(question, _cam_jpeg_b64(frame), "image/jpeg")


def _presence_loop() -> None:
    """LOCAL ambient awareness — greets you on arrival, notices when you leave,
    nudges breaks, and streams a live thumbnail to the HUD. All free."""
    if not _cascades_load():
        print("📷 Webcam presence disabled (OpenCV unavailable).")
        return
    global _paused_for_away, _locked_for_away
    last_thumb = 0.0
    last_posture = time.time()
    _PRESENCE["tick"] = time.time()
    while True:
        if not _PRESENCE["on"]:
            _cam_release()
            time.sleep(2)
            continue
        frame = _cam_frame()
        now = time.time()
        # accumulate at-desk time
        dt = now - _PRESENCE["tick"]
        _PRESENCE["tick"] = now
        if _PRESENCE["present"] and 0 < dt < 10:
            _PRESENCE["desk_seconds"] += dt
        if frame is not None:
            if _face_present(frame):
                _PRESENCE["last_seen"] = now
            is_present = (now - _PRESENCE["last_seen"]) < 8.0   # debounce flicker

            # keep a short rolling buffer of recent frames for "camera recall"
            _CAM_RING.append((now, _cam_jpeg_b64(frame, 320, 45)))
            if len(_CAM_RING) > 40:           # ~2 min at 3s cadence, in-memory only
                _CAM_RING.pop(0)

            if is_present and not _PRESENCE["present"]:          # ── arrived ──
                # If you were away a while and movement appears, optionally ping your phone.
                if _AWAY_ALERT and JARVIS_TG_TOKEN and JARVIS_TG_CHAT and \
                        _PRESENCE["away_since"] and now - _PRESENCE["away_since"] > 300:
                    try:
                        p, err = _cam_photo(os.path.join(tempfile.gettempdir(), "jarvis_alert.jpg"))
                        if p:
                            _tg_send_photo(p, caption="👀 Movement detected at your desk.")
                    except Exception:
                        pass
                _PRESENCE["present"] = True
                _PRESENCE["since"] = now
                away = now - (_PRESENCE["away_since"] or now)
                _hud_emit("presence", present=True)
                # restore things paused while away
                if _paused_for_away:
                    _media_control("play"); _paused_for_away = False
                _locked_for_away = False
                if _PRESENCE["away_since"] and away > 120 and not _ABORT.is_set():
                    mins = max(1, round(away / 60))
                    recap = [t for ts, t in _AWAY_LOG if ts >= _PRESENCE["away_since"]]
                    msg = f"Welcome back. You were away about {mins} minute{'s' if mins != 1 else ''}."
                    if recap:
                        msg += " While you were gone: " + "; ".join(recap[:4]) + "."
                    speak(msg)
                    _AWAY_LOG.clear()
                else:
                    _hud_emit("toast", text="👤 you're here", level="ok")
            elif (not is_present) and _PRESENCE["present"]:      # ── left ──
                _PRESENCE["present"] = False
                _PRESENCE["away_since"] = now
                _PRESENCE["away_count"] += 1
                _hud_emit("presence", present=False)
                if _AWAY_PAUSE:
                    _media_control("pause"); _paused_for_away = True

            if _PRESENCE["present"]:
                # break nudge after ~45 min continuous (local, free)
                if now - _PRESENCE["since"] > 2700 and now - _PRESENCE["last_break"] > 2700:
                    _PRESENCE["last_break"] = now
                    speak("You've been at it for a while. Maybe stretch and hydrate.")
                # optional Gemini posture check
                if _POSTURE_EVERY and now - last_posture > _POSTURE_EVERY * 60:
                    last_posture = now
                    try:
                        verdict = _cam_vision("Look at this side view. Is the person slouching "
                                              "or is their posture good? Reply 'SLOUCHING: <tip>' "
                                              "or 'GOOD'.")
                        if verdict.upper().startswith("SLOUCH"):
                            speak(verdict.split(":", 1)[-1].strip() or "Sit up straight.")
                    except Exception:
                        pass
                # live thumbnail to HUD (~every 4s) — local only
                if now - last_thumb > 4:
                    last_thumb = now
                    _hud_emit("cam", img="data:image/jpeg;base64," + _cam_jpeg_b64(frame, 240, 45),
                              present=True)
            else:
                # locked-on-leave: after the grace period, lock the workstation once
                if _AWAY_LOCK_SEC and not _locked_for_away and \
                        _PRESENCE["away_since"] and now - _PRESENCE["away_since"] > _AWAY_LOCK_SEC:
                    _locked_for_away = True
                    try:
                        import ctypes
                        ctypes.windll.user32.LockWorkStation()
                    except Exception:
                        pass
        time.sleep(3)


def _camera_recall(question: str) -> str:
    """Review the recent rolling webcam frames (in-memory, local) and answer —
    e.g. 'did anyone come to my desk while I was away?'"""
    if not _CAM_RING:
        return "I don't have any recent camera frames buffered yet."
    frames = _CAM_RING[-16:]
    n = len(frames)
    k = min(4, n)
    idxs = [int(i * (n - 1) / (k - 1)) for i in range(k)] if k > 1 else [n - 1]
    b64s = [frames[i][1] for i in idxs]
    span = max(1, int(frames[-1][0] - frames[0][0]))
    return _gemini_vision_multi(
        f"These are {k} webcam frames captured over the last ~{span} seconds (oldest first), "
        f"facing the user's desk. {question}", b64s)


def _focus_report() -> str:
    mins = int(_PRESENCE["desk_seconds"] / 60)
    state = "at your desk" if _PRESENCE.get("present") else "away"
    return (f"This session you've been at your desk about {mins} minute(s), "
            f"and stepped away {_PRESENCE['away_count']} time(s). Right now you're {state}.")


def _cam_photo(path: str = "") -> tuple:
    """Capture a still from the webcam to disk. Returns (path, error)."""
    frame = _cam_frame()
    if frame is None:
        return "", "no webcam frame"
    import cv2
    if not path:
        pics = os.path.join(os.path.expanduser("~"), "Pictures")
        path = os.path.join(pics if os.path.isdir(pics) else tempfile.gettempdir(),
                            "jarvis_photo.jpg")
    else:
        path = os.path.expanduser(path.strip().strip('"'))
    try:
        cv2.imwrite(path, frame)
        return path, ""
    except Exception as e:
        return "", str(e)


def _scan_qr() -> str:
    """Decode a QR code held up to the webcam (local, free via OpenCV)."""
    import cv2
    for _ in range(6):                 # a few tries to let the user position it
        frame = _cam_frame()
        if frame is not None:
            try:
                data, _pts, _ = cv2.QRCodeDetector().detectAndDecode(frame)
                if data:
                    return data
            except Exception:
                pass
        time.sleep(0.4)
    return ""


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  TELEGRAM BRIDGE  ─  command Jarvis from your phone, anywhere
#  Free (Telegram Bot API). Runs the SAME brain + tools as voice, but over
#  chat. Locked to your chat id only. Dependency-free (raw HTTP via requests).
#    Setup: message @BotFather → /newbot → put the token in .env as
#           JARVIS_TELEGRAM_TOKEN=...   then message your bot once; it replies
#           with your chat id → add JARVIS_TELEGRAM_CHAT_ID=... and restart.
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
JARVIS_TG_TOKEN = os.environ.get("JARVIS_TELEGRAM_TOKEN", "").strip()
JARVIS_TG_CHAT  = os.environ.get("JARVIS_TELEGRAM_CHAT_ID", "").strip()
_BRAIN_LOCK = threading.Lock()   # serialize voice vs phone so they never collide
_tg_history = []                 # separate conversation context for the phone channel


def _tg_call(method: str, **params):
    try:
        r = requests.post(f"https://api.telegram.org/bot{JARVIS_TG_TOKEN}/{method}",
                          json=params, timeout=70)
        return r.json()
    except Exception:
        return {}


def _tg_send(text: str, chat_id=None):
    if not text:
        return
    # Telegram caps messages at 4096 chars.
    _tg_call("sendMessage", chat_id=chat_id or JARVIS_TG_CHAT, text=str(text)[:4000])


def _tg_send_photo(path: str, chat_id=None, caption: str = ""):
    try:
        with open(path, "rb") as f:
            requests.post(f"https://api.telegram.org/bot{JARVIS_TG_TOKEN}/sendPhoto",
                          data={"chat_id": chat_id or JARVIS_TG_CHAT, "caption": caption[:1000]},
                          files={"photo": f}, timeout=60)
    except Exception:
        pass


def _notify(text: str, speak_it: bool = True, to_phone: bool = True) -> None:
    """Tell the user something — out loud and (if configured) pushed to their phone."""
    if speak_it:
        try:
            speak(text)
        except Exception:
            pass
    if to_phone and JARVIS_TG_TOKEN and JARVIS_TG_CHAT:
        _tg_send("🔔 " + text)


def _tg_voice_to_text(file_id: str) -> str:
    """Download a Telegram voice note (.oga / Opus), convert to wav via ffmpeg,
    and transcribe with Gemini. Returns the text (or '')."""
    try:
        fp = (_tg_call("getFile", file_id=file_id).get("result") or {}).get("file_path")
        if not fp:
            return ""
        data = requests.get(f"https://api.telegram.org/file/bot{JARVIS_TG_TOKEN}/{fp}", timeout=60).content
        src = tempfile.NamedTemporaryFile(delete=False, suffix=".oga")
        src.write(data); src.close()
        wav = src.name + ".wav"
        try:
            subprocess.run(["ffmpeg", "-y", "-i", src.name, "-ar", "16000", "-ac", "1", wav],
                           capture_output=True, timeout=40)
            with open(wav, "rb") as f:
                b64 = base64.b64encode(f.read()).decode()
            return _gemini_transcribe(b64, "wav")
        finally:
            for p in (src.name, wav):
                try:
                    os.unlink(p)
                except Exception:
                    pass
    except Exception:
        return ""


def _tg_message(text: str, voice, chat_id: str) -> None:
    """Resolve a phone message (transcribing a voice note if needed), then run it."""
    if voice and not text:
        _tg_send("🎙 transcribing…", chat_id)
        text = _tg_voice_to_text(voice.get("file_id", ""))
        if not text:
            _tg_send("Sorry, I couldn't understand that voice note.", chat_id)
            return
        _tg_send(f"🎙 heard: “{text}”", chat_id)
    if text:
        _tg_handle(text, chat_id)


def _tg_handle(text: str, chat_id: str) -> None:
    """Handle one authorized phone message (runs in its own thread)."""
    low = text.lower().strip()
    if low in ("/start", "/help"):
        _tg_send("Jarvis here. Send me any command and I'll run it on your PC.\n"
                 "/screenshot — my screen   /see — webcam   /stop — abort", chat_id)
        return
    if low.startswith("/see") or low.startswith("/webcam"):
        try:
            p, err = _cam_photo(os.path.join(tempfile.gettempdir(), "jarvis_see.jpg"))
            if p:
                _tg_send_photo(p, chat_id, "Live webcam")
                os.unlink(p)
            else:
                _tg_send(f"Couldn't access the webcam: {err}", chat_id)
        except Exception as e:
            _tg_send(f"Webcam error: {e}", chat_id)
        return
    if low.startswith("/stop"):
        _ABORT.set()
        _tg_send("Stopping the current task.", chat_id)
        return
    if low.startswith("/screen"):
        try:
            b64, _mime = _grab_screen_b64()
            p = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg").name
            with open(p, "wb") as f:
                f.write(base64.b64decode(b64))
            _tg_send_photo(p, chat_id, "Your screen right now")
            os.unlink(p)
        except Exception as e:
            _tg_send(f"Couldn't grab the screen: {e}", chat_id)
        return
    # Normal command → run it through the same brain/tools as voice.
    global _tg_history
    _tg_send("On it… 🛠", chat_id)
    _ABORT.clear()
    try:
        with _BRAIN_LOCK:
            _hud_emit("user", text=f"📱 {text}")
            resp, _tg_history = process_command(text, _tg_history)
    except Exception as e:
        resp = f"Something went wrong: {e}"
    _tg_send(resp if resp and resp != "__ABORTED__" else "Okay, stopped.", chat_id)


def _telegram_loop() -> None:
    if not JARVIS_TG_TOKEN:
        return
    me = _tg_call("getMe").get("result", {})
    print(f"📱 Telegram bridge active as @{me.get('username','?')}"
          + ("" if JARVIS_TG_CHAT else "  (message it once to get your chat id)"))
    offset = 0
    while True:
        # server long-poll 50s, HTTP client 70s → comfortable margin so a normal
        # long-poll doesn't time out the request and churn-reconnect.
        resp = _tg_call("getUpdates", offset=offset, timeout=50)
        if not resp.get("ok"):
            time.sleep(3)
            continue
        for upd in resp.get("result", []):
            offset = upd["update_id"] + 1
            msg = upd.get("message") or upd.get("edited_message") or {}
            chat_id = str((msg.get("chat") or {}).get("id", ""))
            text = (msg.get("text") or "").strip()
            voice = msg.get("voice") or msg.get("audio")
            if not chat_id or (not text and not voice):
                continue
            if not JARVIS_TG_CHAT:                       # not yet authorized — help them set it
                _tg_send(f"Your chat id is {chat_id}.\nAdd  JARVIS_TELEGRAM_CHAT_ID={chat_id}  "
                         "to your .env and restart me to authorize this chat.", chat_id)
                continue
            if chat_id != JARVIS_TG_CHAT:                # ignore everyone except you
                continue
            threading.Thread(target=_tg_message, args=(text, voice, chat_id), daemon=True).start()


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  SCHEDULER  ─  run commands later / on a daily schedule, autonomously
#  One-shot ("in 2 hours…") or daily ("every day at 08:00…"). When a task fires
#  Jarvis runs it through the full brain + tools and notifies you (voice + phone).
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
_SCHED_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "jarvis_schedule.json")
_SCHEDULE = []
_SCHED_LOCK = threading.Lock()


def _sched_load():
    global _SCHEDULE
    try:
        with open(_SCHED_PATH, encoding="utf-8") as f:
            _SCHEDULE = json.load(f)
    except Exception:
        _SCHEDULE = []


def _sched_save():
    try:
        with open(_SCHED_PATH, "w", encoding="utf-8") as f:
            json.dump(_SCHEDULE, f)
    except Exception:
        pass


def _schedule_task(command: str, delay_seconds: int = 0, at_time: str = "", daily: bool = False) -> str:
    item = {"command": command}
    if at_time:
        item.update({"time": at_time.strip(), "daily": bool(daily), "last": ""})
        when = f"every day at {at_time}" if daily else f"at {at_time}"
    else:
        item["at"] = time.time() + max(5, int(delay_seconds or 0))
        when = f"in about {max(1, round((delay_seconds or 0) / 60))} minute(s)"
    with _SCHED_LOCK:
        _SCHEDULE.append(item)
        _sched_save()
    _hud_emit("toast", text="🗓 scheduled", level="ok")
    return f"Scheduled: \"{command}\" {when}."


def _schedule_loop():
    while True:
        now = time.time()
        hhmm = datetime.datetime.now().strftime("%H:%M")
        today = datetime.date.today().isoformat()
        due = []
        with _SCHED_LOCK:
            keep = []
            for it in _SCHEDULE:
                if "at" in it:
                    (due if it["at"] <= now else keep).append(it)
                elif it.get("time") == hhmm and it.get("last") != today:
                    it["last"] = today
                    due.append(it); keep.append(it)        # daily → stays
                else:
                    keep.append(it)
            if len(keep) != len(_SCHEDULE) or due:
                _SCHEDULE[:] = keep
                _sched_save()
        for it in due:
            cmd = it.get("command", "")
            try:
                _notify(f"Running your scheduled task: {cmd}", speak_it=False)
                with _BRAIN_LOCK:
                    resp, _ = process_command(cmd, [])
                _notify(f"Scheduled task done — {cmd[:50]}: {str(resp)[:300]}")
            except Exception as e:
                _notify(f"Scheduled task failed ({cmd[:40]}): {e}", speak_it=False)
        time.sleep(15)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  EXTENDED TOOLBOX  ─  files · system · text/AI · web · utilities · fun
#  Local ops are free/instant; web ones use free, no-key APIs; AI ones reuse
#  the Gemini brain. All wrapped so a failure returns a friendly string.
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def _xp(path: str) -> str:
    return os.path.expanduser((path or "").strip().strip('"'))


# ── files & folders ──────────────────────────────────────────────────────────
def _list_files(path: str = "") -> str:
    p = _xp(path) or os.path.expanduser("~")
    if not os.path.isdir(p):
        return f"Not a folder: {p}"
    try:
        items = sorted(os.listdir(p))
    except Exception as e:
        return f"Couldn't list {p}: {e}"
    out = []
    for it in items[:60]:
        full = os.path.join(p, it)
        out.append(f"  {'[dir] ' if os.path.isdir(full) else '      '}{it}")
    extra = f"\n  …and {len(items) - 60} more" if len(items) > 60 else ""
    return f"{p}  ({len(items)} items):\n" + "\n".join(out) + extra


def _search_files(query: str, path: str = "") -> str:
    root = _xp(path) or os.path.expanduser("~")
    q = query.lower()
    hits = []
    for dp, dns, fns in os.walk(root):
        if sum(s in dp.lower() for s in ("\\appdata\\", "\\node_modules\\", "\\.git\\")):
            continue
        for fn in fns:
            if q in fn.lower():
                hits.append(os.path.join(dp, fn))
                if len(hits) >= 40:
                    return f"Found {len(hits)}+ (showing 40):\n" + "\n".join("  " + h for h in hits)
    return (f"Found {len(hits)}:\n" + "\n".join("  " + h for h in hits)) if hits else f"No files matching '{query}' under {root}."


def _open_folder(path: str) -> str:
    p = _xp(path)
    if not os.path.exists(p):
        return f"Path doesn't exist: {p}"
    try:
        os.startfile(p if os.path.isdir(p) else os.path.dirname(p))
        return f"Opened {p} in Explorer."
    except Exception as e:
        return f"Couldn't open: {e}"


def _create_folder(path: str) -> str:
    p = _xp(path)
    try:
        os.makedirs(p, exist_ok=True)
        return f"Created folder {p}."
    except Exception as e:
        return f"Couldn't create folder: {e}"


def _move_path(src: str, dst: str) -> str:
    import shutil
    try:
        shutil.move(_xp(src), _xp(dst))
        return f"Moved to {_xp(dst)}."
    except Exception as e:
        return f"Move failed: {e}"


def _copy_path(src: str, dst: str) -> str:
    import shutil
    s, d = _xp(src), _xp(dst)
    try:
        shutil.copytree(s, d) if os.path.isdir(s) else shutil.copy2(s, d)
        return f"Copied to {d}."
    except Exception as e:
        return f"Copy failed: {e}"


def _delete_path(path: str) -> str:
    """Safe delete — moves to a ~/.jarvis_trash folder (reversible), never hard-deletes."""
    import shutil
    p = _xp(path)
    if not os.path.exists(p):
        return f"Path doesn't exist: {p}"
    trash = os.path.join(os.path.expanduser("~"), ".jarvis_trash")
    os.makedirs(trash, exist_ok=True)
    try:
        dest = os.path.join(trash, os.path.basename(p.rstrip("\\/")))
        if os.path.exists(dest):
            dest += f"_{int(_PRESENCE.get('tick', 0))}"
        shutil.move(p, dest)
        return f"Moved to trash ({trash}). Recover it there if needed."
    except Exception as e:
        return f"Delete failed: {e}"


def _rename_path(path: str, new_name: str) -> str:
    p = _xp(path)
    try:
        dest = os.path.join(os.path.dirname(p), new_name)
        os.rename(p, dest)
        return f"Renamed to {dest}."
    except Exception as e:
        return f"Rename failed: {e}"


def _zip_path(path: str) -> str:
    import shutil
    p = _xp(path)
    try:
        base = p.rstrip("\\/")
        out = shutil.make_archive(base, "zip", p if os.path.isdir(p) else os.path.dirname(p),
                                  None if os.path.isdir(p) else os.path.basename(p))
        return f"Zipped to {out}."
    except Exception as e:
        return f"Zip failed: {e}"


def _unzip_file(path: str, dest: str = "") -> str:
    import shutil
    p = _xp(path)
    d = _xp(dest) or os.path.splitext(p)[0]
    try:
        shutil.unpack_archive(p, d)
        return f"Extracted to {d}."
    except Exception as e:
        return f"Unzip failed: {e}"


def _disk_usage(path: str = "") -> str:
    import shutil
    p = _xp(path) or "C:\\"
    try:
        t, u, f = shutil.disk_usage(p)
        gb = lambda x: round(x / 1024**3, 1)
        return f"{p}: {gb(f)} GB free of {gb(t)} GB ({round(u/t*100)}% used)."
    except Exception as e:
        return f"Disk usage error: {e}"


def _file_info(path: str) -> str:
    p = _xp(path)
    if not os.path.exists(p):
        return f"Path doesn't exist: {p}"
    st = os.stat(p)
    when = datetime.datetime.fromtimestamp(st.st_mtime).strftime("%Y-%m-%d %H:%M")
    size = st.st_size
    unit = "bytes"
    for u in ("KB", "MB", "GB"):
        if size >= 1024:
            size /= 1024; unit = u
    return f"{p}\n  {'folder' if os.path.isdir(p) else 'file'}, {round(size,1)} {unit}, modified {when}"


# ── system ───────────────────────────────────────────────────────────────────
def _system_info() -> str:
    import platform
    parts = [f"{platform.system()} {platform.release()}", f"machine {platform.machine()}"]
    try:
        import psutil
        vm = psutil.virtual_memory()
        parts.append(f"CPU {psutil.cpu_count()} cores @ {round(psutil.cpu_percent())}%")
        parts.append(f"RAM {round(vm.used/1024**3,1)}/{round(vm.total/1024**3,1)} GB")
    except Exception:
        pass
    g = _gpu_stats()
    if g:
        parts.append(f"GPU {round(g['util'])}% , {round(g['temp'])}°C")
    return " | ".join(parts)


def _list_processes(top: int = 10) -> str:
    try:
        import psutil
        procs = []
        for pr in psutil.process_iter(["name", "memory_info"]):
            try:
                procs.append((pr.info["name"], pr.info["memory_info"].rss))
            except Exception:
                pass
        procs.sort(key=lambda x: x[1], reverse=True)
        return "Top processes by memory:\n" + "\n".join(
            f"  {n}  —  {round(m/1024**2)} MB" for n, m in procs[:top])
    except Exception as e:
        return f"Process list error: {e}"


def _kill_process(name: str) -> str:
    try:
        import psutil
        killed = 0
        for pr in psutil.process_iter(["name"]):
            try:
                if name.lower() in (pr.info["name"] or "").lower():
                    pr.terminate(); killed += 1
            except Exception:
                pass
        return f"Closed {killed} process(es) matching '{name}'." if killed else f"No process matching '{name}'."
    except Exception as e:
        return f"Kill error: {e}"


def _battery_status() -> str:
    try:
        import psutil
        b = psutil.sensors_battery()
        if not b:
            return "No battery detected (desktop?)."
        plug = "charging" if b.power_plugged else "on battery"
        left = "" if b.power_plugged or b.secsleft < 0 else f", ~{b.secsleft//3600}h{(b.secsleft%3600)//60}m left"
        return f"Battery {round(b.percent)}% ({plug}{left})."
    except Exception as e:
        return f"Battery error: {e}"


def _screenshot_save(path: str = "") -> str:
    from PIL import ImageGrab
    p = _xp(path) or os.path.join(os.path.expanduser("~"), "Pictures", "jarvis_screen.png")
    try:
        os.makedirs(os.path.dirname(p), exist_ok=True)
        ImageGrab.grab().save(p)
        return f"Saved screenshot to {p}."
    except Exception as e:
        return f"Screenshot failed: {e}"


def _set_clipboard(text: str) -> str:
    try:
        import pyperclip
        pyperclip.copy(text)
        return "Copied to clipboard."
    except Exception as e:
        return f"Clipboard error: {e}"


def _ip_info() -> str:
    try:
        d = requests.get("http://ip-api.com/json", timeout=10).json()
        if d.get("status") != "success":
            return "Couldn't get IP info."
        return (f"Public IP {d.get('query')} — {d.get('city')}, {d.get('regionName')}, "
                f"{d.get('country')} (ISP: {d.get('isp')}).")
    except Exception as e:
        return f"IP info error: {e}"


# ── text / AI utilities ──────────────────────────────────────────────────────
def _ai_text(instruction: str, text: str, temp: float = 0.3) -> str:
    try:
        return (_gemini_request([{"role": "user", "content": f"{instruction}\n\n{text}"}],
                                None, temp).get("content") or "").strip()
    except Exception as e:
        return f"AI error: {e}"


def _generate_password(length: int = 16, symbols: bool = True) -> str:
    import secrets, string
    n = max(6, min(64, int(length or 16)))
    alpha = string.ascii_letters + string.digits + ("!@#$%^&*-_=+" if symbols else "")
    return "Generated password: " + "".join(secrets.choice(alpha) for _ in range(n))


def _hash_text(text: str, algo: str = "sha256") -> str:
    import hashlib
    algo = algo.lower()
    if algo not in hashlib.algorithms_available:
        algo = "sha256"
    return f"{algo}: {hashlib.new(algo, text.encode()).hexdigest()}"


def _base64_tool(text: str, mode: str = "encode") -> str:
    try:
        if mode == "decode":
            return "Decoded: " + base64.b64decode(text).decode("utf-8", "replace")
        return "Encoded: " + base64.b64encode(text.encode()).decode()
    except Exception as e:
        return f"Base64 error: {e}"


def _format_json(text: str) -> str:
    try:
        return "```json\n" + json.dumps(json.loads(text), indent=2) + "\n```"
    except Exception as e:
        return f"Invalid JSON: {e}"


def _count_words(text: str) -> str:
    words = len(text.split())
    return f"{words} words, {len(text)} characters, {len(text.splitlines()) or 1} line(s)."


def _qr_generate(data: str, path: str = "") -> str:
    try:
        import cv2
        p = _xp(path) or os.path.join(os.path.expanduser("~"), "Pictures", "jarvis_qr.png")
        os.makedirs(os.path.dirname(p), exist_ok=True)
        img = cv2.QRCodeEncoder_create().encode(data)
        cv2.imwrite(p, cv2.resize(img, (400, 400), interpolation=cv2.INTER_NEAREST))
        return f"QR code saved to {p}."
    except Exception as e:
        return f"QR generate error: {e}"


def _image_b64(path: str):
    import cv2
    img = cv2.imread(_xp(path))
    if img is None:
        return None
    ok, buf = cv2.imencode(".jpg", img, [cv2.IMWRITE_JPEG_QUALITY, 80])
    return base64.b64encode(buf).decode()


def _describe_image(path: str, question: str = "Describe this image.") -> str:
    b = _image_b64(path)
    if not b:
        return f"Couldn't read image: {path}"
    return _gemini_vision(question, b, "image/jpeg")


# ── geo helper (shared) ──────────────────────────────────────────────────────
def _geocode(location: str):
    try:
        g = requests.get("https://geocoding-api.open-meteo.com/v1/search",
                         params={"name": location, "count": 1}, timeout=10).json()
        r = (g.get("results") or [None])[0]
        return r
    except Exception:
        return None


# ── free web info ────────────────────────────────────────────────────────────
def _get_forecast(location: str, days: int = 3) -> str:
    r = _geocode(location)
    if not r:
        return f"Couldn't find '{location}'."
    try:
        w = requests.get("https://api.open-meteo.com/v1/forecast", params={
            "latitude": r["latitude"], "longitude": r["longitude"],
            "daily": "temperature_2m_max,temperature_2m_min,weather_code,precipitation_probability_max",
            "timezone": "auto", "forecast_days": max(1, min(7, days))}, timeout=10).json()
        d = w["daily"]
        out = [f"Forecast for {r['name']}:"]
        for i, day in enumerate(d["time"]):
            out.append(f"  {day}: {round(d['temperature_2m_min'][i])}–{round(d['temperature_2m_max'][i])}°C, "
                       f"{_WMO.get(d['weather_code'][i],'?')}, {d['precipitation_probability_max'][i]}% rain")
        return "\n".join(out)
    except Exception as e:
        return f"Forecast error: {e}"


def _air_quality(location: str) -> str:
    r = _geocode(location)
    if not r:
        return f"Couldn't find '{location}'."
    try:
        a = requests.get("https://air-quality-api.open-meteo.com/v1/air-quality", params={
            "latitude": r["latitude"], "longitude": r["longitude"],
            "current": "pm2_5,pm10,us_aqi", "timezone": "auto"}, timeout=10).json()["current"]
        aqi = a.get("us_aqi")
        rating = ("good" if aqi <= 50 else "moderate" if aqi <= 100 else "unhealthy for sensitive groups"
                  if aqi <= 150 else "unhealthy" if aqi <= 200 else "very unhealthy") if aqi is not None else "?"
        return f"{r['name']} air quality: AQI {aqi} ({rating}), PM2.5 {a.get('pm2_5')}, PM10 {a.get('pm10')}."
    except Exception as e:
        return f"Air quality error: {e}"


def _sunrise_sunset(location: str) -> str:
    r = _geocode(location)
    if not r:
        return f"Couldn't find '{location}'."
    try:
        s = requests.get("https://api.sunrise-sunset.org/json", params={
            "lat": r["latitude"], "lng": r["longitude"], "formatted": 0,
            "tzid": r.get("timezone", "UTC")}, timeout=10).json()["results"]
        sr = s["sunrise"][11:16]; ss = s["sunset"][11:16]
        return f"{r['name']}: sunrise {sr}, sunset {ss} (local time)."
    except Exception as e:
        return f"Sunrise/sunset error: {e}"


def _hacker_news(count: int = 5) -> str:
    try:
        ids = requests.get("https://hacker-news.firebaseio.com/v0/topstories.json", timeout=10).json()[:max(1, min(10, count))]
        out = []
        for i in ids:
            it = requests.get(f"https://hacker-news.firebaseio.com/v0/item/{i}.json", timeout=10).json()
            out.append(f"  • {it.get('title','')} ({it.get('score',0)} pts)")
        return "Hacker News top stories:\n" + "\n".join(out)
    except Exception as e:
        return f"Hacker News error: {e}"


def _github_repo(repo: str) -> str:
    repo = repo.strip().replace("https://github.com/", "").strip("/")
    try:
        d = requests.get(f"https://api.github.com/repos/{repo}", timeout=10,
                         headers={"Accept": "application/vnd.github+json"}).json()
        if "full_name" not in d:
            return f"Couldn't find repo '{repo}'."
        return (f"{d['full_name']} — {d.get('description','')}\n"
                f"  ⭐ {d.get('stargazers_count')} stars, {d.get('forks_count')} forks, "
                f"{d.get('open_issues_count')} open issues, language: {d.get('language')}.")
    except Exception as e:
        return f"GitHub error: {e}"


def _synonyms(word: str) -> str:
    try:
        d = requests.get("https://api.datamuse.com/words", params={"rel_syn": word, "max": 10}, timeout=10).json()
        syns = [x["word"] for x in d]
        return f"Synonyms for '{word}': " + ", ".join(syns) if syns else f"No synonyms found for '{word}'."
    except Exception as e:
        return f"Thesaurus error: {e}"


def _random_fact() -> str:
    try:
        return "Did you know? " + requests.get("https://uselessfacts.jsph.pl/api/v2/facts/random", timeout=10).json()["text"]
    except Exception as e:
        return f"Fact error: {e}"


def _tell_joke() -> str:
    try:
        j = requests.get("https://official-joke-api.appspot.com/random_joke", timeout=10).json()
        return f"{j['setup']} … {j['punchline']}"
    except Exception as e:
        return f"Joke error: {e}"


def _this_day() -> str:
    try:
        d = requests.get("https://history.muffinlabs.com/date", timeout=10).json()
        events = d["data"]["Events"][:4]
        return f"On this day ({d['date']}):\n" + "\n".join(f"  {e['year']}: {e['text']}" for e in events)
    except Exception as e:
        return f"History error: {e}"


def _stock_price(symbol: str) -> str:
    sym = symbol.upper().strip()
    try:
        d = requests.get(f"https://query1.finance.yahoo.com/v8/finance/chart/{sym}",
                         timeout=10, headers={"User-Agent": "Mozilla/5.0"}).json()
        res = d["chart"]["result"][0]["meta"]
        price = res.get("regularMarketPrice")
        prev = res.get("chartPreviousClose") or res.get("previousClose")
        cur = res.get("currency", "")
        chg = f" ({(price-prev)/prev*100:+.1f}%)" if price and prev else ""
        return f"{sym}: {price} {cur}{chg}."
    except Exception:
        return f"Couldn't get a price for '{symbol}' (try the exact ticker, e.g. AAPL, TSLA)."


def _time_in(city: str) -> str:
    """Local time in a city/timezone using the system tz database (no API)."""
    zones = {"new york": "America/New_York", "london": "Europe/London", "tokyo": "Asia/Tokyo",
             "paris": "Europe/Paris", "dubai": "Asia/Dubai", "sydney": "Australia/Sydney",
             "los angeles": "America/Los_Angeles", "san francisco": "America/Los_Angeles",
             "singapore": "Asia/Singapore", "berlin": "Europe/Berlin", "moscow": "Europe/Moscow",
             "delhi": "Asia/Kolkata", "mumbai": "Asia/Kolkata", "india": "Asia/Kolkata",
             "beijing": "Asia/Shanghai", "hong kong": "Asia/Hong_Kong", "toronto": "America/Toronto"}
    key = city.lower().strip()
    tz = zones.get(key, city if "/" in city else None)
    if not tz:
        return f"I don't know the timezone for '{city}'. Try a major city or an IANA zone like 'Asia/Tokyo'."
    try:
        from zoneinfo import ZoneInfo
        now = datetime.datetime.now(ZoneInfo(tz))
        return f"Time in {city.title()}: {now.strftime('%I:%M %p, %A %d %b')}."
    except Exception as e:
        return f"Time lookup error: {e}"


# ── fun / quick math (local) ─────────────────────────────────────────────────
def _roll_dice(sides: int = 6, count: int = 1) -> str:
    sides = max(2, min(1000, int(sides or 6)))
    count = max(1, min(20, int(count or 1)))
    rolls = [_random.randint(1, sides) for _ in range(count)]
    return f"🎲 Rolled {count}d{sides}: {rolls}" + (f" (total {sum(rolls)})" if count > 1 else "")


def _flip_coin() -> str:
    return "🪙 " + _random.choice(["Heads", "Tails"])


def _random_number(lo: int = 1, hi: int = 100) -> str:
    lo, hi = int(lo), int(hi)
    if lo > hi:
        lo, hi = hi, lo
    return f"🔢 {_random.randint(lo, hi)} (between {lo} and {hi})"


def _days_until(date_str: str) -> str:
    try:
        for fmt in ("%Y-%m-%d", "%d-%m-%Y", "%d/%m/%Y", "%m/%d/%Y"):
            try:
                target = datetime.datetime.strptime(date_str.strip(), fmt).date()
                break
            except ValueError:
                target = None
        if not target:
            return f"Couldn't parse the date '{date_str}'. Try YYYY-MM-DD."
        delta = (target - datetime.date.today()).days
        if delta == 0:
            return "That's today!"
        return f"{abs(delta)} day(s) {'until' if delta > 0 else 'ago'} {target.isoformat()}."
    except Exception as e:
        return f"Date error: {e}"


def _expand_url(url: str) -> str:
    """Resolve a shortened/redirecting URL to its final destination (read-only)."""
    if not url.startswith("http"):
        url = "https://" + url
    try:
        r = requests.head(url, allow_redirects=True, timeout=10)
        if r.url == url:
            r = requests.get(url, allow_redirects=True, timeout=10, stream=True)
        return f"{url}\n  → {r.url}"
    except Exception as e:
        return f"Couldn't expand that URL: {e}"


def _gemini_available() -> bool:
    """True if Gemini is reachable and at least one key authenticates. A 429/503
    still counts as available (working, just busy/limited — runtime rotation
    handles that). Only a network failure or all-keys-rejected means False."""
    if not GEMINI_KEYS:
        return False
    body = {"model": GEMINI_MODEL,
            "messages": [{"role": "user", "content": "hi"}], "max_tokens": 1}
    for key in GEMINI_KEYS:
        try:
            r = requests.post(
                GEMINI_URL,
                headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
                json=body, timeout=15,
            )
        except Exception:
            return False   # network down → Gemini unreachable
        if r.status_code in (200, 429, 503):
            return True
    return False


def wish() -> None:
    hour = datetime.datetime.now().hour
    if hour < 12:
        speak("Good morning! Jarvis online.")
    elif hour < 18:
        speak("Good afternoon! Jarvis ready.")
    else:
        speak("Good evening! Jarvis at your service.")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  MAIN LOOP
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
if __name__ == "__main__":
    # ── Launch the live HUD (web control center) ──────────────────────────────
    _hud_url = _start_hud()
    if _hud_url:
        try:
            webbrowser.open(_hud_url)
        except Exception:
            pass
    _hud_emit("state", state="standby", sub="booting")

    # Warm the voice detector (torch + Silero) in the background so startup stays
    # fast (~2s) while the model loads during the greeting — ready by first listen.
    threading.Thread(target=_get_vad, daemon=True).start()

    _mem_load()   # load long-term memory from disk
    if _MEM:
        print(f"Memory: {len(_MEM)} fact(s) loaded.")
    _hud_emit("memory", count=len(_MEM))

    _usage_load()         # restore today's Gemini usage so the meter survives restarts
    _hud_quota()

    _emit_memory_list()   # populate the HUD memory drawer

    _load_skills()        # load any self-authored skills from jarvis_skills/
    if _SKILLS:
        print(f"Skills: {len(_SKILLS)} self-authored skill(s) loaded — {', '.join(_SKILLS)}.")

    _rem_load()   # restore any pending reminders from a previous session
    threading.Thread(target=_reminder_loop, daemon=True).start()   # fires reminders when due
    threading.Thread(target=_vitals_loop, daemon=True).start()     # streams live system telemetry
    threading.Thread(target=_routine_loop, daemon=True).start()    # proactive daily briefings
    threading.Thread(target=_watch_loop, daemon=True).start()      # proactive screen-watch (off until enabled)
    if JARVIS_TG_TOKEN:
        threading.Thread(target=_telegram_loop, daemon=True).start()   # phone control via Telegram
    _sched_load()
    threading.Thread(target=_schedule_loop, daemon=True).start()       # scheduled automations
    if _SCHEDULE:
        print(f"Scheduler: {len(_SCHEDULE)} task(s) loaded.")
    if _PRESENCE["on"]:
        threading.Thread(target=_presence_loop, daemon=True).start()   # webcam visual presence (local)
        print(f"📷 Visual presence ON (camera index {JARVIS_CAM_INDEX}) — local face detection, "
              "free. Set JARVIS_CAM_INDEX or JARVIS_PRESENCE=0 to change.")

    print("=" * 62)
    print("  JARVIS  —  AI Assistant")
    print(f"  Brain  : Gemini {' → '.join(GEMINI_MODELS)}  ({len(GEMINI_KEYS)} keys, fully cloud — no local GPU)")
    print("  Web    : Playwright Chromium  (DOM-based, reliable)")
    print("  Search : DuckDuckGo  (real-time, free)")
    print("  Apps   : pywinauto   (Windows accessibility API)")
    if _hud_url:
        print(f"  HUD    : {_hud_url}   (live control center — opening in browser)")
    print("  Say 'Jarvis stop' or 'Jarvis goodbye' to exit.")
    print("=" * 62)

    print("Checking AI brain...")
    if not GEMINI_KEYS:
        print("No Gemini keys found. Add them to .env (JARVIS_GEMINI_KEYS=key1,key2) "
              "or jarvis_keys.txt.")
        speak("I have no Gemini keys configured. Please add your keys and restart.")
        exit(1)
    if _gemini_available():
        _hud_emit("online", up=True)
        _hud_emit("brain", provider="gemini", model=GEMINI_MODEL, key=1)
        print(f"Brain ready: Gemini ({len(GEMINI_KEYS)} key(s), {len(GEMINI_MODELS)} models). "
              "Fully cloud — nothing runs on your GPU.\n")
    else:
        # Keys exist but Gemini is unreachable right now (likely no internet).
        _hud_emit("online", up=False)
        print("⚠ Gemini not reachable right now (check your internet). "
              "Jarvis needs the network to think.")
        speak("I can't reach Gemini right now. Please check your internet connection.")

    wish()

    print(_install_abort_hotkey())

    conversation_history: list = []
    WAKE_WORDS = ("jarvis ", "hey jarvis ", "ok jarvis ", "okay jarvis ")

    awaiting_followup = False   # conversational mode: brief window after a reply
    while True:
        # Typed commands from the HUD jump the queue and skip the wake word.
        typed = None
        try:
            typed = _TEXT_CMDS.get_nowait()
        except _queue.Empty:
            pass

        if typed is not None:
            query, is_typed = typed, True
        else:
            query, is_typed = takeCommand(), False

        if not query or query == "none":
            awaiting_followup = False        # silence → go back to requiring the wake word
            continue

        q = query.lower().strip()

        # ── Activation gate ───────────────────────────────────────────────────
        # Normally needs the wake word. But right after Jarvis replies we open a
        # short follow-up window where you can just keep talking, no "Jarvis".
        # Typed commands are always activated.
        followup = (awaiting_followup or is_typed) and ("jarvis" not in q)
        awaiting_followup = False             # this turn consumes the window

        if "jarvis" not in q and not followup and not is_typed:
            print("   (no wake word — ignored)")
            continue

        # ── Hard stop — no LLM round-trip ─────────────────────────────────────
        if any(p in q for p in ["jarvis stop", "jarvis goodbye", "jarvis bye",
                                  "jarvis goodnight", "jarvis exit", "jarvis quit"]):
            speak("Goodbye!")
            _hud_emit("toast", text="Jarvis signing off", level="info")
            _hud_emit("state", state="standby", sub="offline")
            _cam_release()        # free the webcam on the way out
            try:
                if _pw_instance:
                    _pw_instance.stop()
            except Exception:
                pass
            break

        if followup:
            print("   ↪ follow-up (no wake word needed)")
            clean = re.sub(r"(?i)\bjarvis\b[,\s]*", "", query.strip()).strip()
        else:
            print("   ✅ Wake word detected")
            # ── Strip wake word so the LLM gets clean intent ───────────────────
            clean = query.strip()
            for w in WAKE_WORDS:
                if clean.lower().startswith(w):
                    clean = clean[len(w):].strip()
                    break
            else:
                clean = re.sub(r"(?i)\bjarvis\b[,\s]*", "", clean).strip()

        # Nothing left after stripping (e.g. user just said "Jarvis")
        if not clean:
            speak("Yes?")
            awaiting_followup = True
            continue

        print(f"\n📨 Sending to AI: {clean}")
        _hud_emit("user", text=clean)
        _ABORT.clear()        # fresh task — ignore any stale stop signal
        with _BRAIN_LOCK:     # never run a voice task and a phone task at the same time
            response, conversation_history = process_command(clean, conversation_history)

        if response == "__ABORTED__":
            _ABORT.clear()    # consume the stop; ready for the next command
            print("⏹  Stopped.")
            _hud_emit("toast", text="Task stopped", level="warn")
            _hud_emit("state", state="standby", sub="stopped")
            speak("Okay, stopped. What should I do instead?")
            awaiting_followup = True
            time.sleep(0.3)
            continue

        _hud_emit("task")     # one task completed → bump the counter
        speak(response)
        # Durable facts are captured in-turn by the model's `remember` tool (see
        # system prompt) — no separate local-model pass, so the GPU stays free.
        awaiting_followup = True         # open a follow-up window for a natural back-and-forth

        time.sleep(0.3)
